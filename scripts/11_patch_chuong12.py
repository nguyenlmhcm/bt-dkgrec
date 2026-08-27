#!/usr/bin/env python
"""Va Chuong 1 va Chuong 2 cua v12, sinh ra v13.

Vi sao la va tai cho
--------------------
Chuong 2 chua 2 doi tuong cong thuc OMML — cong thuc (2.1) va (2.2) bieu dien
temporal fact. python-docx khong doc va khong ghi duoc chung. Sinh lai chuong
nay dong nghia xoa sach hai cong thuc do. Chuong 1 khong co cong thuc nhung duoc
va cung mot luot de giu mot script cho mot lan chay.

Bay loi duoc sua, theo muc do nang
----------------------------------
1. §2.6 phat bieu ket qua SAI: "BT-DKGRec-GCN gan LightGCN nhung van thap hon
   nhe". Do la so cua v11 (10 epoch, chua hoi tu). Ban dung lai cho BT-DKGRec
   0.032291 so voi LightGCN 0.023736 tren cohort Original — vuot, khong thap
   hon. Ngoai ra Chuong 2 khong nen bao cao ket qua cua Chuong 4.
2. §2.7 co mot cau cut: "Thay vi trinh bay lai cac cong thuc chuan da duoc su
   dung rong rai trong nghien cuu he goi y." — khong co menh de chinh.
3. §2.7 co loi nhan doi: "Nguoi dung co lich su (nguoi dung co lich su)" — vet
   cua mot lan tim-thay the "warm user" an ca vao phan chu thich.
4. §2.6 co cum vo nghia: "cu hon cua so tai lieu chinh cua de tai".
5. Bang 2.2 dinh nghia warm user thieu dieu kien thu hai, nen khong giai thich
   duoc con so 593 nguoi dung danh gia.
6. §1.3.2 liet ke NAM mo hinh; ma tran thuc nghiem hien tai co SAU.
7. §1.4 co hai tieu de Heading 3 khong danh so, lech voi moi tieu de con lai.

Bo sung
-------
* §1.5 "Noi dung va phuong phap nghien cuu" — chuong 1 hien khong noi gi ve ky
  luat thuc nghiem, von la diem manh nhat cua de tai.
* §2.8 dinh vi MBGCN va KHGT, kem hai muc [19] va [20] vao thu muc tham khao.
  Truoc ban va nay, de an phat bieu alpha va lambda "ke thua" tu hai cong trinh
  do ma khong trich dan chung.

Nguyen tac an toan giong 10_patch_chuong3.py: moi ban va gan vao mot neo duy
nhat; neo khop khac 1 lan thi dung; khong bao gio cham vao doan chua OMML; sau
khi va thi dem lai cong thuc, bang, hinh.

Van phong bam theo `docs/VAN_PHONG_DE_AN.md`: giu thuat ngu tieng Anh, gia tri
metric dung dau CHAM va phan tram dung dau PHAY, khong in dam giua doan, khong
dau gach ngang tu tu, cau dai noi bang dau cham phay, trinh bay chu khong
thuyet phuc, ket luan co phong ho.

Chay:
    python scripts/11_patch_chuong12.py
    python scripts/11_patch_chuong12.py --no-highlight    # ban nop
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.utils.logging import get_logger  # noqa: E402

log = get_logger(__name__)

SOURCE = Path("docs/De_an_thac_si_v12.docx")
OMML = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

TOUCHED: list = []
CHANGELOG: list[tuple[str, str, str]] = []


# ── Tim va sua doan ─────────────────────────────────────────────────────


def find(doc: Document, needle: str):
    """Doan DUY NHAT chua `needle`. Dung neu khong co hoac co nhieu hon mot."""
    hits = [p for p in doc.paragraphs if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(
            f"LOI: neo khop {len(hits)} lan, can dung 1.\n  neo: {needle[:80]!r}\n"
            "Ban v12 da doi khac so voi luc doc. Doc lai roi sua neo."
        )
    para = hits[0]
    if para._p.findall(f".//{OMML}oMath"):
        raise SystemExit(f"LOI: doan neo co chua cong thuc, khong duoc ghi de.\n  {needle[:80]!r}")
    return para


def find_heading(doc: Document, text: str, style: str):
    """Tieu de khop CHINH XAC ca noi dung lan style.

    Can rieng ham nay vi `find` dung so khop chuoi con: chuoi "Doi tuong nghien
    cuu" vua la tieu de vua nam trong cau mo dau cua chinh muc do.
    """
    hits = [p for p in doc.paragraphs
            if p.text.strip() == text and p.style.name == style]
    if len(hits) != 1:
        raise SystemExit(f"LOI: tieu de {text!r} ({style}) khop {len(hits)} lan, can dung 1.")
    return hits[0]


def retext(para, text: str) -> None:
    """Thay noi dung doan, giu dinh dang cua run dau tien."""
    for run in para.runs[1:]:
        run._r.getparent().remove(run._r)
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)
    TOUCHED.append(para)


def insert_after(doc: Document, anchor, texts: list[str], style: str = "Normal"):
    """Chen cac doan moi ngay sau `anchor`, dung thu tu da cho."""
    previous = anchor
    made = []
    for text in texts:
        para = doc.add_paragraph(text, style=style)
        if style == "Normal":
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        previous._p.addnext(para._p)
        previous = para
        made.append(para)
        TOUCHED.append(para)
    return made


def cell_containing(doc: Document, needle: str):
    """O bang DUY NHAT chua `needle`."""
    hits = [c for t in doc.tables for r in t.rows for c in r.cells if needle in c.text]
    if len(hits) != 1:
        raise SystemExit(f"LOI: o bang khop {len(hits)} lan, can dung 1: {needle[:60]!r}")
    return hits[0]


def recell(cell, text: str) -> None:
    """Thay noi dung mot o bang, giu dinh dang run dau."""
    para = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    retext(para, text)


# ── Chuong 1 ────────────────────────────────────────────────────────────


def patch_11_boi_canh(doc: Document) -> None:
    """§1.1: neo boi canh vao tai lieu thay vi de trong."""
    CHANGELOG.append((
        '§1.1',
        'Bổ sung trích dẫn cho đoạn bối cảnh (KG trong recommender, temporal KG, nguồn dữ liệu)',
        'Mục bối cảnh trước đây không có trích dẫn nào'))

    para = find(doc, "Knowledge Graph giúp biểu diễn quan hệ giữa khách hàng, sản phẩm")
    retext(para,
        "Knowledge Graph biểu diễn quan hệ giữa khách hàng, sản phẩm, danh mục và "
        "thuộc tính sản phẩm, và các khảo sát về KG-based recommender systems ghi "
        "nhận ba lợi ích chính là giảm sparsity, hỗ trợ cold-start và tăng khả năng "
        "giải thích [1]. Tuy nhiên, một KG tĩnh không phản ánh trực tiếp sự thay đổi "
        "của hành vi theo thời gian; hướng temporal knowledge graph bổ sung yếu tố "
        "thời gian vào biểu diễn fact nhưng thường được đánh giá bằng temporal link "
        "prediction hoặc KG completion thay vì bằng một giao thức gợi ý Top-K [2]. "
        "Vì vậy, đề tài chọn hướng đồ thị tri thức động để mô hình hóa đồng thời "
        "quan hệ ngữ nghĩa của sản phẩm và yếu tố thời gian trong hành vi khách hàng.")

    para = find(doc, 'Trong phạm vi đề tài, "dự báo hành vi khách hàng" được cụ thể hóa')
    retext(para,
        'Trong phạm vi đề tài, "dự báo hành vi khách hàng" được cụ thể hóa thành bài '
        "toán xếp hạng Top-K theo thời gian. Từ lịch sử view, addtocart, transaction "
        "và tri thức liên quan đến sản phẩm, hệ thống tính điểm cho các cặp khách "
        "hàng-sản phẩm và dự báo những sản phẩm có khả năng phát sinh hành vi mục "
        "tiêu trong giai đoạn tiếp theo. Hành vi mục tiêu gồm addtocart hoặc "
        "transaction; view được sử dụng như tín hiệu lịch sử thể hiện mức độ quan "
        "tâm. Ba loại hành vi này được cung cấp trực tiếp trong bộ dữ liệu Retail "
        "Rocket [7].")


def patch_132_sau_mo_hinh(doc: Document) -> None:
    """§1.3.2: ma tran thuc nghiem co sau mo hinh, khong phai nam."""
    CHANGELOG.append((
        '§1.3.2',
        'Danh sách mô hình so sánh: năm mô hình thành sáu, thêm BT-DKGRec-GCN (λ=0.05)',
        'Mô hình thứ sáu đã có trong mọi bảng Chương 4 nhưng chưa được nêu ở Chương 1'))

    para = find(doc, "Đánh giá mô hình theo giao thức Top-K theo thời gian với hệ mô hình")
    retext(para,
        "Đánh giá mô hình theo giao thức Top-K theo thời gian với hệ mô hình so sánh "
        "gồm sáu bậc: Popularity, Recent Popularity, LightGCN, Static KG-GCN, "
        "BT-DKGRec-GCN với λ=0.01 và BT-DKGRec-GCN với λ=0.05; mỗi bậc thêm đúng một "
        "thành phần so với bậc liền trước, nên phần chênh lệch đo được giữa hai bậc "
        "kề nhau quy về được cho riêng thành phần vừa thêm. Static KG-GCN được giữ "
        "như mô hình đối chứng KG tĩnh để tách ảnh hưởng của tri thức sản phẩm khỏi "
        "cơ chế hành vi-thời gian; hai thiết lập λ được giữ song song để phân biệt "
        "giá trị kế thừa từ công trình trước với giá trị dò được trên tập xác thực.")


def patch_14_danh_so(doc: Document) -> None:
    """§1.4: hai tieu de Heading 3 khong danh so, lech voi phan con lai."""
    CHANGELOG.append((
        '§1.4',
        'Đánh số hai tiêu đề "Đối tượng nghiên cứu" và "Phạm vi nghiên cứu" thành 1.4.1 và 1.4.2',
        'Hai tiêu đề này là Heading 3 nhưng không có số, nên vào mục lục không khớp với các mục còn lại'))

    retext(find_heading(doc, "Đối tượng nghiên cứu", "Heading 3"), "1.4.1. Đối tượng nghiên cứu")
    retext(find_heading(doc, "Phạm vi nghiên cứu", "Heading 3"), "1.4.2. Phạm vi nghiên cứu")


def patch_15_phuong_phap(doc: Document) -> None:
    """§1.5 moi: noi dung va phuong phap nghien cuu."""
    CHANGELOG.append((
        '§1.5',
        'Thêm mục "Nội dung và phương pháp nghiên cứu"',
        'Chương 1 trước đây không nêu phương pháp nghiên cứu và không nêu kỷ luật thực nghiệm'))

    anchor = find(doc, "Ứng dụng nguyên mẫu được xây dựng để truy vết lịch sử hành vi")

    heading = doc.add_paragraph("1.5. Nội dung và phương pháp nghiên cứu", style="Heading 2")
    anchor._p.addnext(heading._p)
    TOUCHED.append(heading)

    insert_after(doc, heading, [
        "Đề tài sử dụng ba phương pháp nghiên cứu bổ trợ nhau. Phương pháp nghiên cứu "
        "tài liệu được dùng để xác định vị trí của đề tài giữa ba hướng liên quan là "
        "KG-based recommendation, dynamic graph recommendation và temporal knowledge "
        "graph, từ đó khoanh phạm vi đóng góp ở mức có thể kiểm chứng được. Phương "
        "pháp mô hình hóa được dùng để chuyển dòng sự kiện có timestamp của Retail "
        "Rocket thành một đồ thị tri thức có trọng số cạnh phụ thuộc loại hành vi và "
        "độ mới thời gian. Phương pháp thực nghiệm đối chứng được dùng để đo phần "
        "đóng góp của từng thành phần, trong đó mỗi cặp mô hình liền kề chỉ khác nhau "
        "đúng một biến.",

        "Nội dung nghiên cứu được triển khai theo năm phần, tương ứng với bố cục của "
        "đề án. Chương 1 xác định bối cảnh, mục tiêu và phạm vi. Chương 2 hệ thống "
        "hóa cơ sở lý thuyết và các công trình liên quan, đồng thời xác định khoảng "
        "trống mà đề tài nhắm tới. Chương 3 trình bày thiết kế schema Behavior-Time "
        "Dynamic Knowledge Graph, cơ chế trọng số hành vi-thời gian, mô hình "
        "BT-DKGRec-GCN và quy trình huấn luyện. Chương 4 báo cáo kết quả thực nghiệm "
        "và phân tích. Chương 5 tổng kết đóng góp, nêu hạn chế và hướng phát triển.",

        "Về mặt kỷ luật thực nghiệm, đề tài áp dụng bốn ràng buộc xuyên suốt cho mọi "
        "mô hình trong ma trận so sánh. Thứ nhất, dữ liệu được chia theo trục thời "
        "gian chứ không chia ngẫu nhiên, và các bước tạo ánh xạ định danh, cắt side "
        "information, dựng tập item ứng viên cùng lấy mẫu âm đều chỉ đọc tập huấn "
        "luyện, nên thông tin của giai đoạn cần dự báo không đi ngược vào quá trình "
        "học. Thứ hai, mọi mô hình được huấn luyện tới khi hội tụ bằng cơ chế early "
        "stopping theo tập xác thực với cùng một ngân sách epoch, thay vì dừng ở một "
        "số epoch cố định. Thứ ba, mỗi cấu hình được chạy trên ba seed và kết quả "
        "được báo cáo dưới dạng trung bình kèm độ lệch chuẩn, do các phép toán sparse "
        "trên GPU không cho kết quả trùng khít giữa các lần chạy. Thứ tư, chênh lệch "
        "giữa các mô hình được kiểm định thống kê trước khi được diễn giải, nhằm phân "
        "biệt cải thiện thực với dao động do khởi tạo ngẫu nhiên.",
    ])


# ── Chuong 2 ────────────────────────────────────────────────────────────


def patch_213_warm_user(doc: Document) -> None:
    """Bang 2.2: dinh nghia warm user thieu dieu kien thu hai."""
    CHANGELOG.append((
        'Bảng 2.2',
        'Bổ sung điều kiện thứ hai vào định nghĩa warm user',
        'Định nghĩa cũ chỉ nêu "có lịch sử trong train", không giải thích được vì sao tập đánh giá chỉ còn 593 người dùng'))

    cell = cell_containing(doc, "Warm user có lịch sử trong train")
    recell(cell,
        "Warm user thoả đồng thời hai điều kiện: có ít nhất một tương tác trong "
        "train, và có ít nhất một item mục tiêu trong giai đoạn đang được đánh giá. "
        "Điều kiện thứ nhất bảo đảm mô hình đã học được embedding riêng cho người "
        "dùng đó; điều kiện thứ hai bảo đảm có đáp án để chấm. Cold user không có "
        "tương tác nào trong train nên mô hình cá nhân hóa không có embedding cho họ; "
        "nhóm này được báo cáo riêng, không trộn vào chỉ số của mô hình cá nhân hóa.")


def patch_24_khoang_trong(doc: Document) -> None:
    """§2.4: cau khuyen nghi thanh cau tuong thuat."""
    CHANGELOG.append((
        '§2.4',
        'Chuyển câu "Khoảng trống phù hợp nên được đặt ở..." thành câu tường thuật',
        'Văn phong khoa học trình bày phạm vi đã chọn, không tự khuyến nghị cho chính mình'))

    para = find(doc, "Khoảng trống phù hợp cho đề tài nên được đặt ở phạm vi hẹp")
    retext(para,
        "Khoảng trống mà đề tài nhắm tới được đặt ở phạm vi hẹp và có thể kiểm chứng: "
        "xây dựng một behavior-time Dynamic Knowledge Graph cho Retail Rocket, trong "
        "đó hành vi view, addtocart, transaction cùng item properties, category tree "
        "và timestamp được đưa vào cùng một quy trình dự báo Top-K sản phẩm có khả "
        "năng phát sinh hành vi mục tiêu theo thời gian [1], [2], [7].")


def patch_26_ket_qua_sai(doc: Document) -> None:
    """§2.6: go phat bieu ket qua sai va khong thuoc Chuong 2."""
    CHANGELOG.append((
        '§2.6',
        'Bỏ phát biểu "BT-DKGRec-GCN gần LightGCN nhưng vẫn thấp hơn nhẹ ở Recall@20 và NDCG@20"',
        'Phát biểu này lấy từ số của bản v11 (10 epoch, chưa mô hình nào hội tụ) và đã bị kết quả dựng lại bác bỏ; ngoài ra Chương 2 không nên báo cáo kết quả của Chương 4'))

    para = find(doc, "Vai trò của LightGCN được dùng để đặt mô hình đề xuất")
    retext(para,
        "Trong đề tài, LightGCN giữ vai trò đối chứng graph collaborative filtering "
        "đủ mạnh: nó dùng cùng split, cùng ánh xạ định danh, cùng tập item ứng viên "
        "và cùng hàm mục tiêu với mô hình đề xuất, chỉ khác ở chỗ không nhận side "
        "information và không nhận trọng số hành vi-thời gian. Nhờ vậy, phần chênh "
        "lệch giữa LightGCN và Static KG-GCN quy về được cho tri thức sản phẩm, còn "
        "phần chênh lệch giữa Static KG-GCN và BT-DKGRec-GCN quy về được cho cơ chế "
        "hành vi-thời gian. Kết quả định lượng của từng cặp so sánh được trình bày ở "
        "Chương 4.")


def patch_26_bpr(doc: Document) -> None:
    """§2.6: cum 'cu hon cua so tai lieu chinh' khong co nghia."""
    CHANGELOG.append((
        '§2.6',
        'Sửa câu về BPR: cụm "cũ hơn cửa sổ tài liệu chính của đề tài" không có nghĩa',
        'Lỗi diễn đạt còn sót lại từ bản trước'))

    para = find(doc, "BPR là nền tảng phương pháp cũ hơn")
    retext(para,
        "BPR ra đời trước khoảng thời gian của các công trình đối chiếu chính trong "
        "đề tài, nên được dùng như nguồn cơ sở cho hàm mục tiêu chứ không dùng làm "
        "bằng chứng về tính mới.")


def patch_27_cau_cut(doc: Document) -> None:
    """§2.7: cau cut, khong co menh de chinh."""
    CHANGELOG.append((
        '§2.7',
        'Sửa câu cụt "Thay vì trình bày lại các công thức chuẩn..."',
        'Câu không có mệnh đề chính'))

    # Tim ca hai doan TRUOC khi sua: doan thay the ben duoi cung chua chuoi neo
    # cua doan thua, nen sua truoc thi neo se khop hai lan.
    para = find(doc, "Thay vì trình bày lại các công thức chuẩn")
    thua = find(doc, "Bảng 2.3 nêu vai trò của từng chỉ số trong đúng ngữ cảnh")

    retext(para,
        "Đề tài sử dụng bốn chỉ số phổ biến cho đánh giá Top-K. Các công thức của "
        "bốn chỉ số này đã được dùng rộng rãi trong nghiên cứu hệ gợi ý nên không "
        "được trình bày lại; thay vào đó, Bảng 2.3 nêu vai trò của từng chỉ số trong "
        "đúng ngữ cảnh thực nghiệm của đề tài [8], [9].")

    # Doan cu tro thanh thua sau khi y cua no duoc gop vao doan tren.
    thua._p.getparent().remove(thua._p)


def patch_27_nhan_doi(doc: Document) -> None:
    """§2.7: loi nhan doi trong danh sach rang buoc danh gia."""
    CHANGELOG.append((
        '§2.7 và §4.1.2',
        'Sửa lỗi nhân đôi "Người dùng có lịch sử (người dùng có lịch sử)" ở cả hai chỗ',
        'Vết của một lần tìm-thay thế thuật ngữ ăn cả vào phần chú thích trong ngoặc'))

    para = find(doc, "(người dùng có lịch sử) là nhóm chính cho personalized")
    retext(para,
        "Người dùng có lịch sử trong train, gọi là warm user, là nhóm chính cho "
        "personalized Top-K evaluation.")

    # Cung mot loi o §4.1.2. Chuong 4 se duoc sinh lai sau, nhung de lai mot loi
    # da biet thi te hon la sua ngay.
    para = find(doc, "(người dùng có lịch sử) là visitor có lịch sử trong train")
    retext(para,
        "Warm user là visitor có lịch sử trong train và có target trong giai đoạn "
        "đánh giá. Cold user không có embedding cá nhân hóa từ train nên được báo "
        "riêng hoặc dùng phương án dự phòng theo độ phổ biến, không trộn vào metric "
        "chính của mô hình cá nhân hóa. Trên cohort Original, tập xác thực có 552 "
        "warm user và tập kiểm tra có 593 warm user.")


def patch_28_mbgcn_khgt(doc: Document) -> None:
    """§2.8: dinh vi MBGCN va KHGT, hai nguon ma de an noi la ke thua."""
    CHANGELOG.append((
        '§2.8',
        'Bổ sung đoạn định vị MBGCN và KHGT, kèm hai mục [19] và [20] vào thư mục tham khảo',
        'Đề án phát biểu α và λ kế thừa từ hai công trình này nhưng không trích dẫn chúng; đây là một lỗ hổng trích dẫn hội đồng dễ hỏi'))

    anchor = find(doc, "LightGCN là mô hình đối chứng lọc cộng tác dựa trên đồ thị đủ mạnh")

    insert_after(doc, anchor, [
        "Hai công trình gợi ý đa hành vi có liên quan trực tiếp đến cơ chế được đề "
        "xuất ở Chương 3. MBGCN mô hình hóa nhiều loại hành vi trên một đồ thị hợp "
        "nhất và học mức đóng góp khác nhau của từng loại hành vi vào hành vi mục "
        "tiêu, thay vì coi mọi tương tác như nhau [19]. KHGT sử dụng kiến trúc graph "
        "transformer phân cấp có tri thức bổ trợ, trong đó một cơ chế mã hóa thời "
        "gian được đưa vào để biểu diễn phản ánh được diễn biến của tương tác theo "
        "thời gian [20]. Đề tài kế thừa ý tưởng phân biệt mức độ ý định giữa các loại "
        "hành vi từ hướng thứ nhất và ý tưởng đưa thời gian của tương tác vào biểu "
        "diễn từ hướng thứ hai.",

        "Khác biệt về cách hiện thực cần được nêu rõ để tránh hiểu nhầm về phạm vi "
        "đóng góp. MBGCN học mức đóng góp của từng hành vi như tham số của mô hình, "
        "trong khi đề tài gán ba giá trị cố định cho ba loại hành vi rồi đưa thẳng "
        "vào trọng số cạnh trước khi huấn luyện. KHGT mã hóa thời gian bên trong cơ "
        "chế attention, trong khi đề tài dùng một hàm suy giảm mũ tính sẵn trên cạnh "
        "và không dùng attention. Hai lựa chọn đơn giản hóa này làm giảm khả năng "
        "biểu diễn của mô hình, nhưng đổi lại cho phép giữ kiến trúc lan truyền cố "
        "định giữa các mô hình đối chứng; nhờ đó cặp Static KG-GCN và BT-DKGRec-GCN "
        "chỉ khác nhau ở hàm tính trọng số cạnh, và phép so sánh giữ được tính kiểm "
        "soát.",
    ])


def patch_thu_muc(doc: Document) -> None:
    """Them [19] MBGCN va [20] KHGT vao cuoi thu muc tham khao."""
    anchor = find(doc, "E. Rossi et al., “Temporal Graph Networks for Deep Learning")

    insert_after(doc, anchor, [
        "[19] B. Jin, C. Gao, X. He, D. Jin, and Y. Li, “Multi-behavior Recommendation "
        "with Graph Convolutional Networks,” in Proc. 43rd Int. ACM SIGIR Conf. "
        "Research and Development in Information Retrieval, 2020, pp. 659-668. "
        "DOI: 10.1145/3397271.3401072.",

        "[20] L. Xia, C. Huang, Y. Xu, P. Dai, X. Zhang, H. Yang, J. Pei, and L. Bo, "
        "“Knowledge-Enhanced Hierarchical Graph Transformer Network for Multi-Behavior "
        "Recommendation,” in Proc. AAAI Conf. Artificial Intelligence, vol. 35, no. 5, "
        "2021, pp. 4486-4493. DOI: 10.1609/aaai.v35i5.16576.",
    ])


# ── Danh muc sua doi, kiem ke, to nen ───────────────────────────────────


def append_changelog(doc: Document) -> None:
    """Noi cac dong moi vao dung bang danh muc sua doi da co cua v12."""
    tables = [t for t in doc.tables
              if t.rows and "Mục" in t.rows[0].cells[0].text and len(t.columns) == 3]
    if len(tables) != 1:
        raise SystemExit(f"LOI: tim thay {len(tables)} bang danh muc sua doi, can dung 1.")
    table = tables[0]
    for muc, sua, ly_do in CHANGELOG:
        cells = table.add_row().cells
        for cell, text in zip(cells, (f"[v13] {muc}", sua, ly_do)):
            cell.text = text
        TOUCHED.append(table)


def census(doc: Document) -> dict:
    return {
        "cong thuc": len(doc.element.body.findall(f".//{OMML}oMath")),
        "bang": len(doc.tables),
        "hinh": len(doc.inline_shapes),
    }


def highlight_touched() -> int:
    n = 0
    for item in TOUCHED:
        is_table = hasattr(item, "rows")
        paragraphs = ([para for row in item.rows for cell in row.cells
                       for para in cell.paragraphs] if is_table else [item])
        for para in paragraphs:
            for run in para.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                n += 1
    return n


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--source", type=Path, default=SOURCE)
    ap.add_argument("--out", type=Path, default=Path("docs/De_an_thac_si_v13.docx"))
    ap.add_argument("--no-highlight", action="store_true",
                    help="Khong to nen vang. Dung cho ban nop.")
    args = ap.parse_args()

    if not args.source.exists():
        raise SystemExit(f"LOI: khong thay {args.source}")

    doc = Document(str(args.source))
    before = census(doc)
    log.info("v12: %s", before)

    patch_11_boi_canh(doc)
    patch_132_sau_mo_hinh(doc)
    patch_14_danh_so(doc)
    patch_15_phuong_phap(doc)

    patch_213_warm_user(doc)
    patch_24_khoang_trong(doc)
    patch_26_ket_qua_sai(doc)
    patch_26_bpr(doc)
    patch_27_cau_cut(doc)
    patch_27_nhan_doi(doc)
    patch_28_mbgcn_khgt(doc)
    patch_thu_muc(doc)

    append_changelog(doc)

    after = census(doc)
    log.info("v13: %s", after)
    if after != before:
        raise SystemExit(
            f"LOI: kiem ke lech.\n  truoc: {before}\n  sau  : {after}\n"
            "Ban va chi duoc thay doan van, khong duoc them bot cong thuc, bang hay hinh."
        )

    if not args.no_highlight:
        log.info("To nen vang %d run tren %d vi tri da cham", highlight_touched(), len(TOUCHED))

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out))
    log.info("Da ghi %s — %d ban va", args.out, len(CHANGELOG))


if __name__ == "__main__":
    main()
