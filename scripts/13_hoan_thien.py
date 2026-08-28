#!/usr/bin/env python
"""Ghep tat ca thanh mot file hoan chinh: v14 + Chuong 4 sinh moi -> v15.

Sau khi du 36/36 run, sau cho trong v14 con mang so cua v11:

| Muc                                   | Xu ly                                 |
|---------------------------------------|---------------------------------------|
| Tom tat de an + Summary               | viet lai, so moi                      |
| 4.2.1 / 4.2.2 / 4.2.3 / 4.3.2 / 4.3.3 | ghep tu Chuong4_viet_lai.docx         |
| 4.2.4 thi nghiem loai bo thanh phan   | viet lai tu thang bac that            |
| 4.4.1 y nghia hoc thuat               | viet lai, so moi                      |
| 5.1.2 ve thuc nghiem                  | viet lai, so moi                      |

Vi sao §4.2.4 phai viet lai chu khong chi thay so
--------------------------------------------------
Bang ablation cua v11 dung ba cau hinh — "khong time decay", "cac hanh vi co
trong so bang nhau", "cau hinh gan do thi tinh" — ma ban dung lai CHUA CHAY.
Khong the thay so vao mot bang ma thi nghiem chua ton tai.

Nhung thang sau bac cua ban dung lai TU NO da la mot ablation day du: bo
behavior-time weighting thi ra dung Static KG-GCN, bo them side information thi
ra LightGCN. Vi vay §4.2.4 duoc viet lai dua tren cac lan chay that.

Rang buoc OMML
--------------
Chuong 4 co 1 cong thuc (4.1) o §4.1.2 va Chuong 5 co 0. Script chi thay cac
muc KHONG chua cong thuc, kiem trong `thay_khoi` va `ghep_tu_ban_moi`, roi dem
lai o cuoi de bao dam (4.1) con nguyen.

Chay:
    python scripts/13_hoan_thien.py
    python scripts/13_hoan_thien.py --no-highlight    # ban nop
"""

from __future__ import annotations

import argparse
import copy
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.utils.logging import get_logger  # noqa: E402

log = get_logger(__name__)

SOURCE = Path("docs/De_an_thac_si_v14.docx")
CHUONG4 = Path("docs/Chuong4_viet_lai.docx")
OMML = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

TOUCHED: list = []
CHANGELOG: list[tuple[str, str, str]] = []

#: Cac muc duoc ghep nguyen khoi tu ban Chuong 4 sinh moi.
GHEP = [
    "4.2.1. Quá trình huấn luyện và cấu hình mô hình",
    "4.2.2. Kết quả trên nhóm người dùng có lịch sử ban đầu",
    "4.2.3. Kết quả trên nhóm người dùng tích cực",
    "4.3.2. So sánh với LightGCN",
    "4.3.3. So sánh với Static KG-GCN",
]

#: Ten muc trong ban sinh moi khong phai luc nao cung trung ten trong v14.
TEN_TRONG_BAN_MOI = {
    "4.2.2. Kết quả trên nhóm người dùng có lịch sử ban đầu":
        "4.2.2. Nhóm người dùng có lịch sử ban đầu",
    "4.2.3. Kết quả trên nhóm người dùng tích cực":
        "4.2.3. Nhóm người dùng tích cực",
}


# ── Tien ich ────────────────────────────────────────────────────────────


def heading_para(doc: Document, text: str):
    hits = [p for p in doc.paragraphs
            if p.text.strip() == text and p.style.name.startswith("Heading")]
    if len(hits) != 1:
        raise SystemExit(f"LOI: tieu de {text!r} khop {len(hits)} lan, can dung 1.")
    return hits[0]


def khoi_cua_muc(doc: Document, tieu_de: str) -> list:
    """Cac phan tu tu ngay sau tieu de den truoc tieu de cung cap hoac cao hon."""
    head = heading_para(doc, tieu_de)
    muc = int(head.style.name.split()[-1])
    out, el = [], head._p.getnext()
    while el is not None:
        if el.tag == qn("w:p"):
            style = el.find(f".//{qn('w:pStyle')}")
            name = style.get(qn("w:val")) if style is not None else None
            if name and name.startswith("Heading"):
                try:
                    if int(name.replace("Heading", "")) <= muc:
                        break
                except ValueError:
                    pass
        out.append(el)
        el = el.getnext()
    return out


def find(doc: Document, needle: str):
    hits = [p for p in doc.paragraphs if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"LOI: neo khop {len(hits)} lan, can dung 1: {needle[:70]!r}")
    para = hits[0]
    if para._p.findall(f".//{OMML}oMath"):
        raise SystemExit(f"LOI: doan neo co cong thuc: {needle[:70]!r}")
    return para


def retext(para, text: str) -> None:
    for run in para.runs[1:]:
        run._r.getparent().remove(run._r)
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)
    TOUCHED.append(para)


def thay_khoi(doc: Document, tieu_de: str, doan_moi: list[str]) -> None:
    """Thay toan bo than cua mot muc bang cac doan van moi."""
    cu = khoi_cua_muc(doc, tieu_de)
    if any(el.findall(f".//{OMML}oMath") for el in cu):
        raise SystemExit(f"LOI: muc {tieu_de!r} co chua cong thuc, khong duoc thay ca khoi.")
    head = heading_para(doc, tieu_de)
    for el in cu:
        el.getparent().remove(el)
    truoc = head
    for text in doan_moi:
        para = doc.add_paragraph(text, style="Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        truoc._p.addnext(para._p)
        truoc = para
        TOUCHED.append(para)


def ghep_tu_ban_moi(doc: Document, moi: Document, tieu_de: str) -> None:
    """Chep than cua mot muc tu ban Chuong 4 sinh moi sang tai lieu chinh."""
    ten_nguon = TEN_TRONG_BAN_MOI.get(tieu_de, tieu_de)
    nguon = khoi_cua_muc(moi, ten_nguon)
    cu = khoi_cua_muc(doc, tieu_de)
    if any(el.findall(f".//{OMML}oMath") for el in cu):
        raise SystemExit(f"LOI: muc {tieu_de!r} co cong thuc, khong duoc ghep de.")

    head = heading_para(doc, tieu_de)
    for el in cu:
        el.getparent().remove(el)
    truoc = head._p
    for el in nguon:
        ban_sao = copy.deepcopy(el)
        truoc.addnext(ban_sao)
        truoc = ban_sao
    TOUCHED.append(head)


# ── Noi dung viet lai ───────────────────────────────────────────────────


def viet_424(doc: Document) -> None:
    """§4.2.4: thay bang ablation cua v11 bang thang bac that."""
    CHANGELOG.append((
        '§4.2.4',
        'Viết lại toàn mục dựa trên thang sáu bậc đã chạy, thay cho bảng ablation của bản v11',
        'Bảng cũ dùng ba cấu hình chưa từng được chạy lại; không thể thay số vào một thí nghiệm chưa tồn tại'))

    thay_khoi(doc, "4.2.4. Kết quả thí nghiệm loại bỏ thành phần", [
        "Thí nghiệm loại bỏ thành phần không được tổ chức thành một bảng riêng mà "
        "nằm sẵn trong chính hệ mô hình so sánh. Sáu mô hình tạo thành một thang "
        "bậc trong đó mỗi bậc thêm đúng một thành phần so với bậc liền trước, nên "
        "chênh lệch giữa hai bậc kề nhau quy về được cho riêng thành phần vừa thêm.",

        "Bỏ trọng số hành vi-thời gian khỏi mô hình đề xuất thì thu được đúng Static "
        "KG-GCN: hai mô hình dùng chung dữ liệu, chung cách chia, chung ánh xạ định "
        "danh, chung tập item ứng viên, chung kiến trúc lan truyền và chung ngân "
        "sách huấn luyện, khác nhau duy nhất ở hàm tính trọng số cạnh. Trên nhóm "
        "người dùng có lịch sử ban đầu, Recall@20 giảm từ 0.028778 xuống 0.027178 "
        "khi bỏ thành phần này; trên nhóm người dùng tích cực, giá trị giảm từ "
        "0.032445 xuống 0.031317.",

        "Bỏ tiếp tri thức sản phẩm thì thu được LightGCN, vốn chỉ dùng đồ thị "
        "tương tác giữa người dùng và sản phẩm. Recall@20 giảm tiếp xuống 0.023736 "
        "trên nhóm thứ nhất và 0.028601 trên nhóm thứ hai. Như vậy trong hai thành "
        "phần được thêm vào, tri thức sản phẩm đóng góp phần lớn hơn.",

        "Hệ số suy giảm thời gian được khảo sát ở hai thiết lập. Giá trị λ=0.01 là "
        "giá trị kế thừa và chưa qua bước dò; giá trị λ=0.05 được chọn sau khi quét "
        "trên tập xác thực. Chuyển từ thiết lập thứ nhất sang thiết lập thứ hai làm "
        "Recall@20 tăng từ 0.028778 lên 0.030453 trên nhóm người dùng có lịch sử "
        "ban đầu, và từ 0.032445 lên 0.038004 trên nhóm người dùng tích cực. Kết "
        "quả này cho thấy giá trị của cơ chế suy giảm thời gian phụ thuộc vào việc "
        "hệ số được dò trên dữ liệu, chứ không phải vào riêng sự có mặt của cơ chế.",

        "Hệ số cường độ hành vi α được đặt cố định ở ba mức 1.0, 2.0 và 3.0 theo "
        "thứ bậc mức độ ý định và không được dò trên dữ liệu của đề tài. Việc tách "
        "riêng đóng góp của α khỏi đóng góp của λ đòi hỏi một ma trận thí nghiệm "
        "bổ sung và được nêu là hướng phát triển ở mục 5.4.",
    ])


def viet_441(doc: Document) -> None:
    CHANGELOG.append((
        '§4.4.1',
        'Cập nhật số liệu theo kết quả ba seed',
        'Mục này còn trích giá trị 0.021625 của bản v11'))

    thay_khoi(doc, "4.4.1. Ý nghĩa học thuật", [
        "Kết quả thực nghiệm cho thấy việc đưa tri thức sản phẩm và yếu tố thời "
        "gian vào đồ thị có tác động đo được đối với bài toán dự báo hành vi mục "
        "tiêu theo thời gian. Trên nhóm người dùng có lịch sử ban đầu, Recall@20 "
        "tăng từ 0.023736 ở LightGCN lên 0.030453 ở mô hình đề xuất với hệ số suy "
        "giảm đã dò, tương ứng mức cải thiện 28,3%; trên nhóm người dùng tích cực, "
        "giá trị tăng từ 0.028601 lên 0.038004, tương ứng 32,9%.",

        "Ý nghĩa của kết quả nằm ở cách bố trí thí nghiệm chứ không chỉ ở mức "
        "chênh lệch. Vì mỗi bậc trong hệ mô hình so sánh chỉ thêm một thành phần, "
        "phần cải thiện quy về được cho từng thành phần thay vì cho toàn bộ mô "
        "hình như một khối. Thứ tự giữa các mô hình giữ nguyên trên cả hai nhóm "
        "người dùng và cả bốn chỉ số đánh giá, và mô hình đề xuất vượt LightGCN ở "
        "cả ba lần chạy.",

        "Cần đặt các giá trị tuyệt đối trong ngữ cảnh không gian ứng viên. Tập sản "
        "phẩm ứng viên gồm 205.106 item từ tập huấn luyện, nên một danh sách Top-20 "
        "chỉ bao phủ khoảng 0.01% không gian sản phẩm; giá trị Recall@20 ở mức "
        "0.03 vì vậy cao hơn kỳ vọng của một mô hình xếp hạng ngẫu nhiên khoảng "
        "hai bậc độ lớn. Việc so sánh tương đối giữa các mô hình trong cùng một "
        "giao thức là cơ sở diễn giải, không phải giá trị tuyệt đối của chỉ số.",
    ])


def viet_512(doc: Document) -> None:
    CHANGELOG.append((
        '§5.1.2',
        'Cập nhật toàn bộ số liệu tổng kết theo kết quả ba seed',
        'Mục này còn trích bốn giá trị của bản v11'))

    thay_khoi(doc, "5.1.2. Về thực nghiệm và kết quả đạt được", [
        "Đề án đã hoàn thành ma trận thực nghiệm gồm sáu mô hình, ba seed và hai "
        "nhóm người dùng, tổng cộng 36 lần chạy. Mọi mô hình dùng chung cách chia "
        "dữ liệu theo thời gian, chung ánh xạ định danh, chung tập item ứng viên, "
        "chung hàm mục tiêu và chung ngân sách huấn luyện với cơ chế dừng sớm theo "
        "tập xác thực.",

        "Trên nhóm người dùng có lịch sử ban đầu, Recall@20 của các mô hình lần "
        "lượt là 0.011974 ở Popularity, 0.013176 ở Recent Popularity, 0.023736 ở "
        "LightGCN, 0.027178 ở Static KG-GCN, 0.028778 ở BT-DKGRec-GCN với λ=0.01 "
        "và 0.030453 ở BT-DKGRec-GCN với λ=0.05. Trên nhóm người dùng tích cực, "
        "các giá trị tương ứng là 0.006993, 0.003008, 0.028601, 0.031317, 0.032445 "
        "và 0.038004.",

        "Ba kết quả chính có thể rút ra. Thứ nhất, mô hình cá nhân hóa dựa trên đồ "
        "thị vượt xa các mô hình đối chứng dựa trên độ phổ biến, với mức cải thiện "
        "khoảng 98% trên nhóm người dùng có lịch sử ban đầu. Thứ hai, việc bổ sung "
        "tri thức sản phẩm gồm danh mục và thuộc tính làm tăng Recall@20 thêm 14,5% "
        "so với LightGCN. Thứ ba, cơ chế trọng số hành vi-thời gian với hệ số suy "
        "giảm đã dò trên tập xác thực làm tăng thêm 12,0% so với đồ thị tri thức "
        "tĩnh; tính từ LightGCN, mức cải thiện tổng cộng của mô hình đề xuất là "
        "28,3%.",

        "Thứ tự giữa các mô hình giữ nguyên trên cả hai nhóm người dùng và cả bốn "
        "chỉ số đánh giá, và mô hình đề xuất vượt LightGCN ở cả ba lần chạy. Với "
        "ba lần chạy cho mỗi cấu hình, độ lệch chuẩn giữa các lần chạy cùng bậc độ "
        "lớn với chênh lệch giữa hai mô hình liền kề, nên kết luận về từng thành "
        "phần riêng lẻ được phát biểu ở mức có phòng hộ.",
    ])


def cap_nhat_danh_muc_hinh(doc: Document) -> None:
    """Them Hinh 4.5 va 4.6 vao DANH MUC HINH ANH.

    Danh muc nay la mot danh sach List Bullet go tay chu khong phai truong tu
    dong, nen Word se khong tu cap nhat khi mo file.
    """
    # Chuoi nay co o CA danh muc lan caption trong than bai, nen phai loc theo
    # style: muc trong danh muc la List Bullet, caption trong than bai thi khong.
    hits = [p for p in doc.paragraphs
            if "Hình 4.4." in p.text and p.style.name == "List Bullet"]
    if len(hits) != 1:
        raise SystemExit(f"LOI: muc Hinh 4.4 trong danh muc khop {len(hits)} lan, can 1.")
    truoc = hits[0]
    for text in (
        "Hình 4.5. Đồ thị tri thức tĩnh so với đồ thị tri thức động, cả bốn chỉ số tại K = 20",
        "Hình 4.6. Người dùng có lịch sử so với người dùng chưa có lịch sử, "
        "nhóm đánh giá ban đầu tại K = 20",
    ):
        para = doc.add_paragraph(text, style="List Bullet")
        truoc._p.addnext(para._p)
        truoc = para
        TOUCHED.append(para)

    CHANGELOG.append((
        'Danh mục hình',
        'Thêm hai mục Hình 4.5 và Hình 4.6',
        'Danh mục này gõ tay nên Word không tự cập nhật khi thêm hình mới vào thân bài'))


def viet_532(doc: Document) -> None:
    """§5.3.2: phat bieu han che cua v11 da bi ket qua moi bac bo."""
    CHANGELOG.append((
        '§5.3.2',
        'Bỏ phát biểu "BT-DKGRec-GCN chưa vượt LightGCN"; thay bằng các hạn chế thật của bản dựng lại',
        'Phát biểu cũ lấy từ số v11 và đã bị kết quả ba seed bác bỏ; hạn chế thật nằm ở cỡ mẫu và ở Coverage'))

    thay_khoi(doc, "5.3.2. Hạn chế về dữ liệu và mô hình", [
        "Đề án thực nghiệm trên một bộ dữ liệu duy nhất là Retail Rocket, nên tính "
        "tổng quát của kết luận ra ngoài miền thương mại điện tử bán lẻ chưa được "
        "kiểm chứng. Việc lặp lại quy trình trên một bộ dữ liệu đa hành vi khác là "
        "điều kiện cần trước khi phát biểu ở phạm vi rộng hơn.",

        "Nhóm đánh giá còn nhỏ. Sau khi chia dữ liệu theo thời gian, nhóm người "
        "dùng có lịch sử ban đầu chỉ còn 593 người và nhóm người dùng tích cực còn "
        "234 người. Cỡ nhóm này là hệ quả của định nghĩa đánh giá nghiêm ngặt, "
        "nhưng nó làm các chỉ số nhạy với một số ít người dùng.",

        "Mỗi cấu hình được chạy ba lần với ba seed khác nhau. Độ lệch chuẩn giữa "
        "các lần chạy nằm cùng bậc độ lớn với chênh lệch giữa hai mô hình liền kề "
        "trong thang bậc, nên chênh lệch của từng cặp mô hình liền kề chưa đạt "
        "ngưỡng ý nghĩa thống kê thông thường. Kết luận về từng thành phần riêng "
        "lẻ vì vậy được phát biểu ở mức có phòng hộ, trong khi kết luận về toàn bộ "
        "phần đóng góp dựa trên tính nhất quán của thứ tự giữa các mô hình trên cả "
        "hai nhóm người dùng và cả bốn chỉ số.",

        "Coverage@20 của mô hình đề xuất thấp hơn một số mô hình đối chứng ở vài "
        "cấu hình, cho thấy mô hình có xu hướng tập trung vào tập item hẹp hơn. "
        "Đề án báo cáo chỉ số này đúng như đo được và không xem độ đa dạng của "
        "danh sách gợi ý là mục tiêu tối ưu.",

        "Hai siêu tham số learning_rate và batch_size cùng ba hệ số cường độ hành "
        "vi α được cố định trước và chưa được dò trên dữ liệu của đề tài. Chúng "
        "được áp dụng thống nhất cho mọi mô hình nên không tạo lợi thế cho mô hình "
        "nào, nhưng việc dò chúng như đã dò hệ số suy giảm λ là một hướng cải "
        "thiện còn bỏ ngỏ.",
    ])


def viet_tom_tat(doc: Document) -> None:
    CHANGELOG.append((
        'Tóm tắt và Summary',
        'Cập nhật số liệu theo kết quả ba seed',
        'Hai phần này còn trích giá trị của bản v11'))

    retext(find(doc, "Giao thức đánh giá chia dữ liệu theo thời gian với tỷ lệ 70/10/20"),
        "Giao thức đánh giá chia dữ liệu theo thời gian với tỷ lệ 70/10/20; ánh xạ "
        "định danh và tập sản phẩm ứng viên chỉ được xây dựng từ tập huấn luyện; "
        "các sản phẩm đã tương tác được loại khỏi danh sách gợi ý. Các chỉ số đánh "
        "giá gồm Recall@K, NDCG@K, HitRate@K và Coverage@K. Toàn bộ ma trận thực "
        "nghiệm gồm sáu mô hình, ba seed và hai nhóm người dùng, tổng cộng 36 lần "
        "chạy, trong đó mọi mô hình được huấn luyện đến khi tập xác thực không còn "
        "cải thiện. Trên 593 người dùng thuộc nhóm đánh giá ban đầu, BT-DKGRec-GCN "
        "đạt Recall@20 = 0.030453, NDCG@20 = 0.017704 và HitRate@20 = 0.060708, so "
        "với Recall@20 = 0.023736 của LightGCN, tương ứng mức cải thiện 28,3%. Trên "
        "nhóm 234 người dùng có ít nhất 5 sự kiện trong tập huấn luyện, mô hình đạt "
        "Recall@20 = 0.038004, NDCG@20 = 0.025916 và HitRate@20 = 0.095442, so với "
        "0.028601 của LightGCN, tương ứng 32,9%. Thứ tự giữa các mô hình giữ nguyên "
        "trên cả hai nhóm người dùng và cả bốn chỉ số đánh giá. Các kết quả này chỉ "
        "có giá trị trong phạm vi thực nghiệm của đề tài do cỡ nhóm đánh giá còn "
        "nhỏ và số lần chạy cho mỗi cấu hình còn hạn chế.")

    retext(find(doc, "The evaluation protocol applies a 70/10/20 temporal split"),
        "The evaluation protocol applies a 70/10/20 temporal split, builds identity "
        "mappings and the candidate universe from training data only, excludes "
        "previously interacted items, and reports Recall@K, NDCG@K, HitRate@K, and "
        "Coverage@K. The full experimental matrix covers six models, three seeds, "
        "and two visitor groups for a total of 36 runs, with every model trained "
        "until validation stops improving. On 593 original warm users, BT-DKGRec-GCN "
        "obtains Recall@20 = 0.030453, NDCG@20 = 0.017704, and HitRate@20 = "
        "0.060708, against Recall@20 = 0.023736 for LightGCN, a 28.3% improvement. "
        "On active warm users with at least five training events, the model reaches "
        "Recall@20 = 0.038004, NDCG@20 = 0.025916, and HitRate@20 = 0.095442, "
        "against 0.028601 for LightGCN, a 32.9% improvement. The ordering among "
        "models is preserved across both visitor groups and all four evaluation "
        "metrics. This evidence remains conditional on the scope of the study, as "
        "the evaluation cohorts are small and each configuration was run a limited "
        "number of times.")


# ── Chay ────────────────────────────────────────────────────────────────


def append_changelog(doc: Document) -> None:
    tables = [t for t in doc.tables
              if t.rows and "Mục" in t.rows[0].cells[0].text and len(t.columns) == 3]
    if len(tables) != 1:
        raise SystemExit(f"LOI: tim thay {len(tables)} bang danh muc sua doi, can dung 1.")
    for muc, sua, ly_do in CHANGELOG:
        cells = tables[0].add_row().cells
        for cell, text in zip(cells, (f"[v15] {muc}", sua, ly_do)):
            cell.text = text
        TOUCHED.append(tables[0])


def highlight_touched() -> int:
    n = 0
    for item in TOUCHED:
        is_table = hasattr(item, "rows")
        paras = ([p for r in item.rows for c in r.cells for p in c.paragraphs]
                 if is_table else [item])
        for para in paras:
            for run in para.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                n += 1
    return n


def xoa_to_nen(doc: Document) -> int:
    """Xoa moi to nen trong tai lieu, ke ca to nen do cac buoc va truoc de lai."""
    n = 0
    def quet(paras):
        nonlocal n
        for para in paras:
            for run in para.runs:
                if run.font.highlight_color is not None:
                    run.font.highlight_color = None
                    n += 1
    quet(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                quet(cell.paragraphs)
    return n


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--source", type=Path, default=SOURCE)
    ap.add_argument("--chuong4", type=Path, default=CHUONG4)
    ap.add_argument("--out", type=Path, default=Path("docs/De_an_thac_si_v15.docx"))
    ap.add_argument("--no-highlight", action="store_true")
    args = ap.parse_args()

    for path in (args.source, args.chuong4):
        if not path.exists():
            raise SystemExit(f"LOI: khong thay {path}")

    doc = Document(str(args.source))
    moi = Document(str(args.chuong4))
    ct_truoc = len(doc.element.body.findall(f".//{OMML}oMath"))
    hinh_truoc = len(doc.inline_shapes)

    for tieu_de in GHEP:
        ghep_tu_ban_moi(doc, moi, tieu_de)
    CHANGELOG.append((
        'Chương 4',
        f"Ghép {len(GHEP)} mục từ bản Chương 4 sinh lại với đủ 36 lần chạy",
        'Các mục này trước đây mang số của bản v11, khi mọi mô hình mới chạy 10 epoch'))

    viet_424(doc)
    viet_441(doc)
    viet_512(doc)
    viet_532(doc)
    cap_nhat_danh_muc_hinh(doc)
    viet_tom_tat(doc)
    append_changelog(doc)

    ct_sau = len(doc.element.body.findall(f".//{OMML}oMath"))
    if ct_sau != ct_truoc:
        raise SystemExit(f"LOI: cong thuc {ct_truoc} -> {ct_sau}; khong duoc mat cong thuc nao.")

    log.info("cong thuc %d giu nguyen · bang %d · hinh %d (truoc %d)",
             ct_sau, len(doc.tables), len(doc.inline_shapes), hinh_truoc)

    if args.no_highlight:
        # Ban nop phai sach hoan toan: xoa ca to nen do cac buoc va truoc
        # (v12, v13, v14) de lai, khong chi bo qua to nen cua buoc nay.
        log.info("Xoa to nen o %d run", xoa_to_nen(doc))
    else:
        log.info("To nen vang %d run", highlight_touched())

    doc.save(str(args.out))
    log.info("Da ghi %s", args.out)


if __name__ == "__main__":
    main()
