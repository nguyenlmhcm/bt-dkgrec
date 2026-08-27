#!/usr/bin/env python
"""Cat bot cong thuc Chuong 3 cua v13, danh so lai, sinh ra v14.

Vi sao cat
----------
Chuong 3 co 33 doi tuong cong thuc tren 31 so hieu. Doc ky thi thay ba loai thua
ro rang, khong phai cat tuy tien:

A. Input/Output cua thuat toan bi danh so nhu cong thuc — (3.11), (3.12),
   (3.19), (3.20). Day la dac ta dau vao/dau ra, thuoc ve pseudocode.
B. Trung lap — (3.13), (3.14), (3.15) chinh la delta_t, w, W(u,i), lap lai
   nguyen (3.16), (3.17), (3.18) o §3.4; (3.22) diem tich vo huong lap lai
   (3.28). Giu ban o muc so huu khai niem.
C. Ky hieu tap hop dan trai o §3.3.2 — (3.2), (3.4), (3.5), (3.7), (3.9). Sau
   cong thuc lien tiep chi de noi G = (V, R, E, A); rieng (3.4) va (3.7) thi
   doan van ngay sau da giai thich lai bang loi.

Tong 13 cong thuc bi cat. Ca tai lieu tu 36 xuong 23. Ba cong thuc dong gop
rieng (3.16)-(3.18) GIU NGUYEN, sau khi danh so lai thanh (3.6)-(3.8) — nam som
hon nen de chi khi bao ve.

Vi sao danh so lai tu dong duoc
-------------------------------
python-docx khong doc va khong ghi duoc than cong thuc OMML, nhung SO HIEU thi
la text thuong nam trong o thu ba cua bang, nen sua duoc. Toan bai chi co ba
tham chieu cheo trong than van — (3.16), (3.18), (3.30) — deu duoc anh xa lai.

Xoa mot cong thuc = xoa ca bang chua no, ke ca doi tuong OMML ben trong. Do la
phep xoa chu khong phai phep sua, nen python-docx lam duoc.

Chay:
    python scripts/12_cat_cong_thuc.py
    python scripts/12_cat_cong_thuc.py --no-highlight    # ban nop
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.utils.logging import get_logger  # noqa: E402

log = get_logger(__name__)

SOURCE = Path("docs/De_an_thac_si_v13.docx")
OMML = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
WT = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"

#: Cac so hieu bi cat, kem loai thua tuong ung.
CAT = {
    "3.2": "C", "3.4": "C", "3.5": "C", "3.7": "C", "3.9": "C",
    "3.11": "A", "3.12": "A", "3.19": "A", "3.20": "A",
    "3.13": "B", "3.14": "B", "3.15": "B", "3.22": "B",
}

TOUCHED: list = []
CHANGELOG: list[tuple[str, str, str]] = []


def find(doc: Document, needle: str):
    hits = [p for p in doc.paragraphs if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(f"LOI: neo khop {len(hits)} lan, can dung 1: {needle[:70]!r}")
    para = hits[0]
    if para._p.findall(f".//{OMML}oMath"):
        raise SystemExit(f"LOI: doan neo co cong thuc, khong duoc ghi de: {needle[:70]!r}")
    return para


def retext(para, text: str) -> None:
    for run in para.runs[1:]:
        run._r.getparent().remove(run._r)
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)
    TOUCHED.append(para)


def drop(doc: Document, needle: str) -> None:
    """Xoa han mot doan van da tro thanh thua."""
    para = find(doc, needle)
    para._p.getparent().remove(para._p)


def bang_cong_thuc(doc: Document):
    """Cac bang co chua OMML, kem danh sach so hieu doc duoc trong bang do."""
    out = []
    for el in doc.element.body:
        if el.tag == qn("w:tbl") and el.findall(f".//{OMML}oMath"):
            text = "".join(x.text or "" for x in el.findall(f".//{WT}"))
            out.append((el, re.findall(r"\((\d+\.\d+)\)", text)))
    return out


# ── Buoc 1: sua van truoc, khi cong thuc con nguyen vi tri ──────────────


def sua_van(doc: Document) -> None:
    """Viet lai cac doan gioi thieu cong thuc sap bi cat, de khong con cau lung lo."""

    # §3.3.2 — (3.2), (3.4), (3.5), (3.7), (3.9)
    retext(find(doc, "Gọi tập visitor, item, category và property lần lượt là"),
        "Đồ thị gồm bốn loại node: visitor, item, category và property. Tập node "
        "của graph là hợp của bốn tập đó:")
    drop(doc, "Tập node của graph được xác định bởi")

    # "tau" con sot lai tu mot lan chuyen ma; ky hieu dung la τ.
    retext(find(doc, "tập E tại thời điểm tau chứa các cạnh"),
        "tập E tại thời điểm τ chứa các cạnh xuất hiện hoặc còn hiệu lực tại "
        "snapshot đó;")

    drop(doc, "Tập loại quan hệ gồm bốn quan hệ chính")
    retext(find(doc, "Trong đó, các ký hiệu quan hệ lần lượt biểu diễn"),
        "Tập loại quan hệ gồm bốn quan hệ chính, lần lượt biểu diễn "
        "INTERACTED_WITH, HAS_CATEGORY, HAS_PROPERTY và PARENT_CATEGORY trong đồ "
        "thị huấn luyện. Tại snapshot thời gian τ, Behavior-Time Dynamic Knowledge "
        "Graph là một đồ thị có hướng, có nhãn quan hệ và có thuộc tính thời gian; "
        "tập cạnh tại snapshot đó gồm cạnh hành vi hợp lệ đến thời điểm τ và các "
        "cạnh tri thức sản phẩm:")
    drop(doc, "Tập cạnh tại snapshot tau gồm cạnh hành vi hợp lệ")
    drop(doc, "Tập thuộc tính của graph được xác định bởi")

    retext(find(doc, "Mỗi dòng trong events.csv thuộc tập train được xem là một bản ghi raw"),
        "Mỗi dòng trong events.csv thuộc tập train là một bản ghi gồm timestamp, "
        "visitor, loại hành vi, item và transaction id nếu bản ghi thuộc giao dịch. "
        "Hàm ánh xạ cạnh chuyển bản ghi raw này thành một fact động:")
    drop(doc, "Trong đó, t là timestamp, u là visitor, b là loại hành vi")

    # §3.3.3 — (3.13), (3.14), (3.15) trung voi §3.4
    retext(find(doc, "Tính tuổi của tương tác theo số ngày"),
        "Tính tuổi của tương tác theo số ngày so với mốc kết thúc train, rồi tính "
        "trọng số của tương tác từ cường độ hành vi và độ mới thời gian; hai đại "
        "lượng này được định nghĩa ở mục 3.4.")
    drop(doc, "Tính trọng số của tương tác")
    retext(find(doc, "Nếu một visitor tương tác nhiều lần với cùng item trong train"),
        "Nếu một visitor tương tác nhiều lần với cùng item trong train, các fact "
        "được tổng hợp theo cặp visitor-item bằng phép cộng trọng số ở mục 3.4.")

    # §3.5.1 — (3.22) trung voi (3.28) o §3.5.2
    retext(find(doc, "Với mỗi item ứng viên, tính điểm phù hợp bằng tích vô hướng"),
        "Với mỗi item ứng viên, tính điểm phù hợp bằng tích vô hướng giữa embedding "
        "visitor và embedding item theo công thức nêu ở mục 3.5.2.")

    CHANGELOG.append((
        'Chương 3',
        'Viết lại các đoạn dẫn vào công thức bị cắt để không còn câu lửng',
        'Xoá công thức mà giữ nguyên câu dẫn sẽ để lại dấu hai chấm không có gì theo sau'))


# ── Buoc 2: xoa bang cong thuc ──────────────────────────────────────────


def xoa_cong_thuc(doc: Document) -> int:
    da_xoa = 0
    for el, so_hieu in bang_cong_thuc(doc):
        if set(so_hieu) & set(CAT):
            if not set(so_hieu) <= set(CAT):
                raise SystemExit(
                    f"LOI: bang {so_hieu} vua co cong thuc phai cat vua co cong thuc "
                    "phai giu. Khong xoa duoc mot phan cua bang.")
            el.getparent().remove(el)
            da_xoa += len(so_hieu)

    for nhom, ten in (("A", "Input/Output của thuật toán bị đánh số như công thức"),
                      ("B", "công thức trùng lặp giữa mục thuật toán và mục mô hình"),
                      ("C", "ký hiệu tập hợp dàn trải ở mục 3.3.2")):
        so = sorted((s for s, g in CAT.items() if g == nhom),
                    key=lambda s: [int(i) for i in s.split(".")])
        CHANGELOG.append((
            'Chương 3',
            f"Cắt {len(so)} công thức ({', '.join(f'({s})' for s in so)}): {ten}",
            'Giảm số công thức để phần đóng góp riêng nổi lên; nội dung được diễn giải bằng lời'))
    return da_xoa


# ── Buoc 3: danh so lai va sua tham chieu cheo ──────────────────────────


def danh_so_lai(doc: Document) -> dict[str, str]:
    """Danh lai so hieu Chuong 3 theo thu tu xuat hien, tra ve anh xa cu -> moi."""
    anh_xa: dict[str, str] = {}
    dem = 0
    can_ghi = []
    for el, so_hieu in bang_cong_thuc(doc):
        if not so_hieu or not so_hieu[0].startswith("3."):
            continue
        for cu in so_hieu:
            dem += 1
            anh_xa[cu] = f"3.{dem}"
        can_ghi.append(el)

    # Ghi thanh MOT luot sau khi da tinh xong toan bo anh xa; ghi xen ke se lam
    # mot so hieu vua doi bi doi them lan nua.
    for el in can_ghi:
        for node in el.findall(f".//{WT}"):
            if node.text and (m := re.fullmatch(r"\s*\((\d+\.\d+)\)\s*", node.text)):
                if (moi := anh_xa.get(m.group(1))):
                    node.text = node.text.replace(m.group(1), moi)

    CHANGELOG.append((
        'Chương 3',
        f"Đánh số lại {dem} công thức còn lại; (3.16)-(3.18) trở thành "
        f"({anh_xa['3.16']})-({anh_xa['3.18']})",
        'Ba công thức đóng góp riêng nằm sớm hơn trong chương nên dễ chỉ khi bảo vệ'))
    return anh_xa


def sua_tham_chieu(doc: Document, anh_xa: dict[str, str]) -> int:
    """Sua cac tham chieu cheo trong than van theo anh xa moi."""
    n = 0
    for para in doc.paragraphs:
        found = re.findall(r"\((\d+\.\d+)\)", para.text)
        if not any(s in anh_xa and anh_xa[s] != s for s in found):
            continue
        moi = re.sub(r"\((\d+\.\d+)\)",
                     lambda m: f"({anh_xa.get(m.group(1), m.group(1))})", para.text)
        retext(para, moi)
        n += 1
    return n


def kiem_tham_chieu_treo(doc: Document, con_lai: set[str]) -> None:
    """Khong duoc con cho nao tro toi cong thuc da bi cat."""
    treo = []
    for para in doc.paragraphs:
        for s in re.findall(r"\((3\.\d+)\)", para.text):
            if s not in con_lai:
                treo.append((s, para.text[:70]))
    if treo:
        raise SystemExit("LOI: tham chieu treo toi cong thuc khong con:\n" +
                         "\n".join(f"  ({s}) trong {t!r}" for s, t in treo))


# ── Chay ────────────────────────────────────────────────────────────────


def append_changelog(doc: Document) -> None:
    tables = [t for t in doc.tables
              if t.rows and "Mục" in t.rows[0].cells[0].text and len(t.columns) == 3]
    if len(tables) != 1:
        raise SystemExit(f"LOI: tim thay {len(tables)} bang danh muc sua doi, can dung 1.")
    for muc, sua, ly_do in CHANGELOG:
        cells = tables[0].add_row().cells
        for cell, text in zip(cells, (f"[v14] {muc}", sua, ly_do)):
            cell.text = text
        TOUCHED.append(tables[0])


def highlight_touched() -> int:
    n = 0
    for item in TOUCHED:
        is_table = hasattr(item, "rows")
        paras = ([p for r in item.rows for c in r.cells for p in c.paragraphs]
                 if is_table else [item])
        for para in paras:
            for run in para.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                n += 1
    return n


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--source", type=Path, default=SOURCE)
    ap.add_argument("--out", type=Path, default=Path("docs/De_an_thac_si_v14.docx"))
    ap.add_argument("--no-highlight", action="store_true")
    args = ap.parse_args()

    if not args.source.exists():
        raise SystemExit(f"LOI: khong thay {args.source}")

    doc = Document(str(args.source))
    truoc = len(doc.element.body.findall(f".//{OMML}oMath"))
    bang_truoc, hinh_truoc = len(doc.tables), len(doc.inline_shapes)

    sua_van(doc)
    da_xoa = xoa_cong_thuc(doc)
    anh_xa = danh_so_lai(doc)
    n_ref = sua_tham_chieu(doc, anh_xa)
    kiem_tham_chieu_treo(doc, set(anh_xa.values()))
    append_changelog(doc)

    sau = len(doc.element.body.findall(f".//{OMML}oMath"))
    if sau != truoc - da_xoa:
        raise SystemExit(f"LOI: cho doi {truoc - da_xoa} cong thuc, dem duoc {sau}.")
    if len(doc.inline_shapes) != hinh_truoc:
        raise SystemExit("LOI: so hinh bi thay doi.")
    if len(doc.tables) != bang_truoc - len({s for s in CAT}):
        raise SystemExit(
            f"LOI: cho doi {bang_truoc - len(CAT)} bang, dem duoc {len(doc.tables)}.")

    log.info("Cong thuc %d -> %d (cat %d) · sua %d tham chieu cheo",
             truoc, sau, da_xoa, n_ref)
    log.info("(3.16)-(3.18) -> (%s)-(%s)", anh_xa["3.16"], anh_xa["3.18"])

    if not args.no_highlight:
        log.info("To nen vang %d run", highlight_touched())

    doc.save(str(args.out))
    log.info("Da ghi %s", args.out)


if __name__ == "__main__":
    main()
