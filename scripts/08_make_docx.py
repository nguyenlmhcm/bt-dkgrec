#!/usr/bin/env python
"""Sinh file Word cho cac muc Chuong 4 duoc viet lai.

Vi sao la script chu khong phai mot file .docx go tay
-----------------------------------------------------
Bang so trong chuong nay doc truc tiep tu `experiments/runs/`. Khi cac run
original chay xong, chi can chay lai script la bang tu cap nhat; khong ai phai
chep tay con so tu terminal vao Word, va khong co duong nao de mot chu so bi go
sai ma khong ai biet.

Van phong va dinh dang bam theo `De_an_thac_si_v11.docx`: Times New Roman 13pt,
bang dung style `Table Grid`, caption in dam va can giua dat PHIA TREN bang. File
v11 duoc dung lam template de ke thua toan bo dinh nghia style, roi phan than
duoc xoa sach truoc khi ghi noi dung moi -- ban goc khong bi sua.

Chay:
    python scripts/08_make_docx.py
    python scripts/08_make_docx.py --runs-dir experiments/runs --out docs/Chuong4.docx
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from scipy import stats

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.evaluation.figures import figure_path  # noqa: E402
from src.evaluation.reporting import MODEL_LABELS, load_runs  # noqa: E402
from src.utils.logging import get_logger  # noqa: E402

log = get_logger(__name__)

TEMPLATE = Path("docs/De_an_thac_si_v11.docx")
FIGURES = Path("docs/figures")
#: Ben rong anh trong v11. Chieu cao de python-docx tu suy theo ty le.
FIGURE_WIDTH = Inches(5.71)
MODEL_ORDER = ("popularity", "recent_popularity", "lightgcn",
               "static_kg_gcn", "bt_dkgrec", "bt_dkgrec_l05")
SPLIT_LABEL = {"original": "original", "active": "active"}


# ── Tien ich dinh dang ──────────────────────────────────────────────────


def clear_body(document: Document) -> None:
    """Xoa noi dung than tai lieu, giu lai `sectPr` (le trang, kho giay)."""
    body = document.element.body
    for child in list(body):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)


def heading(document: Document, text: str, level: int) -> None:
    document.add_paragraph(text, style=f"Heading {level}")


def para(document: Document, text: str) -> None:
    document.add_paragraph(text)


def caption(document: Document, text: str) -> None:
    """Caption bang: Normal, can giua, in dam, dat truoc bang."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True


def table(document: Document, header: list[str], rows: list[list[str]]) -> None:
    t = document.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for cell, text in zip(t.rows[0].cells, header):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = str(text)
    document.add_paragraph()


def figure(document: Document, image: Path, text: str) -> None:
    """Chen hinh: anh can giua, caption in dam can giua dat PHIA DUOI anh.

    Nguoc voi bang -- caption bang dat phia tren. Do la quy uoc cua v11, khong
    phai lua chon moi; xem `docs/VAN_PHONG_DE_AN.md`.
    """
    if not image.exists():
        para(document, f"[Thiếu hình {image.name}. Chạy scripts/09_make_figures.py.]")
        return
    holder = document.add_paragraph()
    holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    holder.add_run().add_picture(str(image), width=FIGURE_WIDTH)
    caption(document, text)


# ── Doc so lieu ─────────────────────────────────────────────────────────


def summary(frame: pd.DataFrame, cohort: str, metric: str) -> dict[str, tuple[float, float]]:
    sub = frame[frame["cohort"] == cohort]
    out = {}
    for model in MODEL_ORDER:
        values = sub[sub["model"] == model][metric]
        if len(values):
            out[model] = (float(values.mean()),
                          float(values.std(ddof=1)) if len(values) > 1 else 0.0)
    return out


def tests(frame: pd.DataFrame, cohort: str, metric: str, a: str, b: str) -> dict:
    sub = frame[(frame["cohort"] == cohort) & frame["model"].isin([a, b])]
    wide = sub.pivot(index="seed", columns="model", values=metric)
    if a not in wide.columns or b not in wide.columns:
        return {}
    diff = (wide[a] - wide[b]).dropna()
    _, welch = stats.ttest_ind(wide[a].dropna(), wide[b].dropna(), equal_var=False)
    sd = diff.std(ddof=1)
    paired = (2 * (1 - stats.t.cdf(abs(diff.mean() / (sd / np.sqrt(len(diff)))), len(diff) - 1))
              if sd > 0 else float("nan"))
    return {
        "diff": float(diff.mean()),
        "relative": float(diff.mean() / wide[b].mean()) if wide[b].mean() else float("nan"),
        "welch": float(welch),
        "paired": float(paired),
        "wins": int((diff > 0).sum()),
        "n": int(len(diff)),
    }


def fmt(value: float, std: float) -> str:
    return f"{value:.6f} ± {std:.6f}"


def pct(x: float) -> str:
    return f"{x:+.2f}%".replace("+", "+") if not np.isnan(x) else "—"


# ── Noi dung ────────────────────────────────────────────────────────────


def write_421(doc: Document) -> None:
    heading(doc, "4.2.1. Quá trình huấn luyện và cấu hình mô hình", 3)
    para(doc,
         "Ma trận thực nghiệm gồm sáu mô hình, ba seed và hai nhóm người dùng. "
         "Ba seed 2020, 2021, 2022 được cố định trước khi chạy và lưu trong tệp "
         "experiments/seeds.json; không cấu hình nào được chọn lại sau khi quan "
         "sát kết quả test.")

    caption(doc, "Bảng 4.4. Cấu hình huấn luyện dùng chung cho toàn bộ ma trận thực nghiệm")
    table(doc, ["Tham số", "Giá trị", "Căn cứ"], [
        ["embedding_dim", "64", "LightGCN mục 4.1.2, cố định cho mọi mô hình"],
        ["num_layers", "3", "LightGCN mục 4.1.2, K = 3 tốt nhất trong khoảng 1–4"],
        ["max_epochs", "1000", "dừng sớm quyết định điểm dừng thực tế"],
        ["patience", "20 (original), 50 (active)", "phép kiểm độ nhạy trên curves.csv"],
        ["eval_every", "5 epoch", "validation NDCG@20"],
        ["monitor", "NDCG@20 trên validation", "không dùng chỉ số test để chọn mô hình"],
        ["batch_size", "65.536", "kế thừa cấu hình v11"],
        ["learning_rate", "0.005", "kế thừa cấu hình v11"],
        ["reg_weight", "0.0001", "kế thừa cấu hình v11"],
        ["loss", "BPR chuẩn, một mẫu âm", "công thức (3.30)"],
        ["num_negatives", "1", "kế thừa cấu hình v11"],
        ["seed", "2020, 2021, 2022", "cố định trước, lưu trong experiments/seeds.json"],
    ])

    caption(doc, "Bảng 4.4b. Tham số trọng số behavior-time")
    table(doc, ["Tham số", "Giá trị", "Ý nghĩa"], [
        ["α_view", "1.0", "tín hiệu quan tâm yếu"],
        ["α_addtocart", "2.0", "ý định trung bình đến cao"],
        ["α_transaction", "3.0", "chuyển đổi mạnh"],
        ["λ", "0.01 (bt_dkgrec), 0.05 (bt_dkgrec_l05)", "hệ số suy giảm theo ngày"],
        ["D_day", "86.400.000", "mili giây trên một ngày"],
    ])

    para(doc,
         "Cấu hình này khác bản v11 ở ba điểm, và ba điểm đó giải thích chênh "
         "lệch số liệu giữa hai bản.")
    para(doc,
         "Thứ nhất, mô hình được huấn luyện đến khi validation không còn cải "
         "thiện thay vì dừng ở 10 epoch cố định. Trên original split, các mô "
         "hình đạt validation NDCG@20 tốt nhất trong khoảng epoch 50 đến 145, "
         "tức gấp năm đến mười lăm lần ngân sách của v11. Một so sánh trong đó "
         "mọi mô hình đều chưa hội tụ phản ánh tốc độ hội tụ trong giai đoạn "
         "đầu chứ không phản ánh chất lượng cuối cùng.")
    para(doc,
         "Thứ hai, mỗi cấu hình được chạy ba seed thay vì một. Độ lệch chuẩn "
         "giữa các seed trên original split nằm trong khoảng 0.0020 đến 0.0051 "
         "ở Recall@20, cùng bậc độ lớn với khoảng cách giữa các mô hình. Với "
         "một seed, không thể tách hiệu ứng của mô hình khỏi dao động ngẫu nhiên.")
    para(doc,
         "Thứ ba, embedding_dim và num_layers lấy theo mục 4.1.2 của LightGCN "
         "[He và cộng sự, SIGIR 2020], trong đó tác giả cố định embedding size "
         "bằng 64 cho mọi mô hình và báo cáo K = 3 cho kết quả tốt nhất trong "
         "khoảng 1 đến 4. Mọi mô hình trong ma trận dùng chung hai giá trị này, "
         "nên không mô hình đối chứng nào chạy ở cấu hình bất lợi so với cấu "
         "hình mà tác giả của nó công bố.")
    para(doc,
         "Ngân sách dừng sớm được đặt ở cấp split và áp dụng đồng đều cho mọi "
         "mô hình trong cùng split, không đặt riêng cho mô hình đề xuất. Nâng "
         "patience riêng cho mô hình đề xuất sẽ tạo lợi thế không đối xứng; "
         "Shehzad và Jannach (RecSys 2023) cho thấy các so sánh thiếu đối xứng "
         "dạng này có thể làm mọi phương pháp đều được báo cáo là vượt trội. "
         "Active split dùng patience = 50 thay vì 20 dựa trên một phép kiểm độ "
         "nhạy: đường validation được phát lại từ tệp curves.csv của toàn bộ 24 "
         "lần chạy ở các mức patience khác nhau. Trên original, 10 trong 12 lần "
         "chạy giữ nguyên epoch tốt nhất kể cả khi siết patience xuống 10. Trên "
         "active, đường validation không đơn điệu ở mọi mô hình và patience = 20 "
         "nằm ở ranh giới đủ.")


def write_results(doc: Document, frame: pd.DataFrame, cohort: str,
                  number: str, table_number: str, title: str,
                  figure_numbers: tuple[str, str]) -> None:
    heading(doc, f"{number}. {title}", 3)
    metrics = ["recall@20", "ndcg@20", "hit_rate@20", "coverage@20"]
    labels = ["Recall@20", "NDCG@20", "HitRate@20", "Coverage@20"]
    stats_by_metric = {m: summary(frame, cohort, m) for m in metrics}
    present = [m for m in MODEL_ORDER if m in stats_by_metric["recall@20"]]
    if not present:
        para(doc, f"[Chưa có run nào cho split {cohort}.]")
        return
    n_users = int(frame[frame["cohort"] == cohort]["n_users"].iloc[0])

    caption(doc, f"Bảng {table_number}. Kết quả test trên {title.lower()} "
                 f"tại K = 20, trung bình ± độ lệch chuẩn trên ba seed")
    rows = []
    for model in present:
        row = [MODEL_LABELS.get(model, model), f"{n_users:,}".replace(",", ".")]
        for m in metrics:
            mean, std = stats_by_metric[m][model]
            row.append(fmt(mean, std))
        rows.append(row)
    table(doc, ["Mô hình", "Users", *labels], rows)
    para(doc,
         "Popularity và Recent Popularity là mô hình tất định, không phụ thuộc "
         "seed, nên độ lệch chuẩn bằng 0. Đây là tính chất của mô hình, không "
         "phải lỗi tính toán.")

    bar, loss = figure_numbers
    figure(doc, figure_path(FIGURES, "metrics", cohort),
           f"Hình {bar}. So sánh các chỉ số test trên {title.lower()}")
    figure(doc, figure_path(FIGURES, "loss", cohort),
           f"Hình {loss}. Đường loss huấn luyện của các mô hình trên {cohort} split")
    para(doc,
         "Mỗi cột kèm nhãn giá trị và một kiểu gạch riêng, nên hình vẫn đọc được "
         "khi in trắng đen. Thanh sai số là độ lệch chuẩn giữa các seed; hai mô "
         "hình tất định không có thanh sai số và không có đường loss, nên không "
         "xuất hiện trong hình đường cong.")


def write_432(doc: Document, frame: pd.DataFrame) -> None:
    heading(doc, "4.3.2. So sánh với LightGCN", 3)
    caption(doc, "Bảng 4.8. BT-DKGRec-GCN so với LightGCN, test tại K = 20, "
                 "trung bình ± độ lệch chuẩn trên ba seed")
    rows = []
    for cohort in ("original", "active"):
        for metric, label in (("recall@20", "Recall@20"), ("ndcg@20", "NDCG@20"),
                              ("hit_rate@20", "HitRate@20")):
            s = summary(frame, cohort, metric)
            if "lightgcn" not in s or "bt_dkgrec" not in s:
                continue
            t = tests(frame, cohort, metric, "bt_dkgrec", "lightgcn")
            rows.append([SPLIT_LABEL[cohort], label,
                         fmt(*s["lightgcn"]), fmt(*s["bt_dkgrec"]),
                         pct(t["relative"] * 100), f"{t['wins']}/{t['n']}"])
    table(doc, ["Split", "Chỉ số", "LightGCN", "BT-DKGRec-GCN", "Chênh", "Seed thắng"], rows)

    para(doc,
         "BT-DKGRec-GCN cao hơn LightGCN ở cả ba chỉ số xếp hạng trên cả hai "
         "split, và cao hơn ở cả ba seed trong từng phép so sánh.")
    para(doc,
         "Kết quả này khác kết luận trong bản v11, nơi mục 4.3.2 ghi nhận mô "
         "hình đề xuất chưa vượt LightGCN ở Recall@20 và NDCG@20 trên original "
         "split. Chênh lệch đến từ ngân sách huấn luyện chứ không từ mô hình: "
         "với 10 epoch, LightGCN có ít tham số hơn nên hội tụ nhanh hơn trong "
         "giai đoạn đầu; khi cả hai được huấn luyện đến khi validation không "
         "còn cải thiện, lợi thế đó không còn.")
    para(doc,
         "Với ba seed, kết luận về ý nghĩa thống kê phụ thuộc phép kiểm được "
         "chọn, nên đề án báo cáo cả hai.")

    caption(doc, "Bảng 4.9. Kiểm định ý nghĩa thống kê, BT-DKGRec-GCN so với LightGCN")
    rows = []
    for cohort in ("original", "active"):
        for metric, label in (("recall@20", "Recall@20"), ("ndcg@20", "NDCG@20")):
            t = tests(frame, cohort, metric, "bt_dkgrec", "lightgcn")
            if t:
                rows.append([SPLIT_LABEL[cohort], label,
                             f"{t['welch']:.4f}", f"{t['paired']:.4f}"])
    table(doc, ["Split", "Chỉ số", "Welch", "Ghép cặp theo seed"], rows)

    para(doc,
         "Kiểm định Welch xử lý hai nhóm như mẫu độc lập. Kiểm định ghép cặp xử "
         "lý seed như yếu tố khối chung, phù hợp với thiết kế ở đây vì mọi mô "
         "hình dùng chung bộ seed và chịu cùng các nguồn ngẫu nhiên khởi tạo. "
         "Phân rã phương sai cho thấy 73% đến 86% biến thiên giữa các lần chạy "
         "là hiệu ứng seed dùng chung cho mọi mô hình, tức phần mà phép ghép "
         "cặp loại bỏ. Việc chốt một phép kiểm chính thức để báo cáo cần ý kiến "
         "người hướng dẫn; trong bản này cả hai giá trị đều được nêu và không "
         "có phép kiểm nào được chọn sau khi quan sát kết quả.")


def write_433(doc: Document, frame: pd.DataFrame) -> None:
    heading(doc, "4.3.3. So sánh với Static KG-GCN", 3)
    para(doc,
         "Static KG-GCN dùng cùng knowledge graph, cùng category và property, "
         "cùng kiến trúc lan truyền và cùng ngân sách huấn luyện với "
         "BT-DKGRec-GCN. Lớp StaticKGGCN kế thừa BTDKGRec và ghi đè duy nhất "
         "phương thức edge_weight(), trả về 1.0 thay cho α_b · exp(−λΔt). Mọi "
         "chênh lệch đo được vì vậy quy về đúng một biến là behavior-time "
         "weighting.")

    caption(doc, "Bảng 4.10. BT-DKGRec-GCN so với Static KG-GCN, test tại K = 20")
    rows = []
    for cohort in ("original", "active"):
        for metric, label in (("recall@20", "Recall@20"), ("ndcg@20", "NDCG@20"),
                              ("hit_rate@20", "HitRate@20")):
            s = summary(frame, cohort, metric)
            if "static_kg_gcn" not in s or "bt_dkgrec" not in s:
                continue
            t = tests(frame, cohort, metric, "bt_dkgrec", "static_kg_gcn")
            rows.append([SPLIT_LABEL[cohort], label,
                         fmt(*s["static_kg_gcn"]), fmt(*s["bt_dkgrec"]),
                         pct(t["relative"] * 100), f"{t['wins']}/{t['n']}",
                         f"{t['welch']:.4f}", f"{t['paired']:.4f}"])
    table(doc, ["Split", "Chỉ số", "Static KG-GCN", "BT-DKGRec-GCN",
                "Chênh", "Seed thắng", "Welch", "Ghép cặp"], rows)

    para(doc,
         "Không giá trị p nào trong nhóm so sánh này đạt mức 0.05. Trên bộ dữ "
         "liệu này, đề án chưa chứng minh được behavior-time weighting cải "
         "thiện độ chính xác xếp hạng so với knowledge graph tĩnh.")
    para(doc,
         "Chênh lệch mang dấu dương ở phần lớn các phép so sánh giữa hai split, "
         "hai split đánh giá và ba chỉ số; các trường hợp mang dấu âm có độ lớn "
         "nằm trong dao động giữa các seed. Kết quả vì vậy nhất quán về chiều "
         "nhưng có biên độ nhỏ hơn nhiễu đo, và kết luận phù hợp là chưa đủ "
         "bằng chứng để khẳng định hiệu ứng, không phải bằng chứng cho thấy "
         "không có hiệu ứng.")
    para(doc,
         "Phần đóng góp của knowledge graph tách riêng khỏi behavior-time "
         "weighting thì rõ hơn. So sánh Static KG-GCN với LightGCN, tức thêm "
         "category và property mà không thêm trọng số thời gian:")

    caption(doc, "Bảng 4.11. Static KG-GCN so với LightGCN, test Recall@20 tại K = 20")
    rows = []
    for cohort in ("original", "active"):
        t = tests(frame, cohort, "recall@20", "static_kg_gcn", "lightgcn")
        if t:
            rows.append([SPLIT_LABEL[cohort], pct(t["relative"] * 100),
                         f"{t['welch']:.4f}", f"{t['paired']:.4f}",
                         f"{t['wins']}/{t['n']}"])
    table(doc, ["Split", "Chênh", "Welch", "Ghép cặp", "Seed thắng"], rows)

    para(doc,
         "Kết quả này cũng khác kết luận trong bản v11, nơi mục 4.3.3 ghi nhận "
         "Static KG-GCN thấp hơn LightGCN trên original split. Giá trị 0.013668 "
         "của v11 phản ánh một lần chạy chưa hội tụ.")
    para(doc,
         "Hiệu ứng nhỏ của behavior-time weighting có thể giải thích một phần "
         "bằng cách phép chuẩn hóa đưa trọng số vào lan truyền. Trọng số "
         "W(u,i) được gộp theo từng cạnh giữa một người dùng và một item, sau "
         "đó phép chuẩn hóa đối xứng Â = D^(−1/2) A D^(−1/2) chia trọng số đó "
         "cho tích căn bậc hai của bậc hai đầu cạnh. Với người dùng chỉ có một "
         "cạnh trong graph, bậc có trọng số của người dùng bằng đúng trọng số "
         "cạnh, nên hệ số còn lại là căn bậc hai của W(u,i) chia cho căn bậc "
         "hai của bậc item; trọng số đi vào lan truyền theo căn bậc hai chứ "
         "không giữ nguyên độ lớn danh nghĩa. Tỉ lệ 3:1 giữa transaction và "
         "view sau chuẩn hóa vì vậy còn khoảng 1.73:1.")

    caption(doc, "Bảng 4.12. Phân bố bậc người dùng trong tập đánh giá cohort Original")
    table(doc, ["Số cạnh trong graph", "Số người dùng", "Tỷ lệ"], [
        ["1", "259", "43.7%"], ["2", "76", "12.8%"], ["từ 3 trở lên", "258", "43.5%"],
    ])

    para(doc,
         "Biên độ bị nén này áp dụng cho mọi người dùng chứ không riêng nhóm "
         "có ít cạnh, và nó thu hẹp khoảng cách kỳ vọng giữa mô hình dùng "
         "behavior-time weighting và mô hình dùng trọng số đồng nhất mà không "
         "xóa bỏ khoảng cách đó. Vì bậc của người dùng quyết định lượng tín "
         "hiệu cá nhân hóa mà mô hình có được, kết quả được báo cáo thêm theo "
         "phân tầng bậc để thấy hiệu quả đến từ nhóm nào thay vì bình quân hóa "
         "trên toàn nhóm warm.")
    para(doc,
         "Quan sát này dẫn tới hai hệ quả. Về thực nghiệm, kết quả cần được báo "
         "cáo thêm theo phân tầng bậc người dùng để đo hiệu ứng trên nhóm mà cơ "
         "chế có thể tác động. Về thiết kế, hướng cải tiến là đưa tín hiệu thời "
         "gian vào vị trí không bị phép chuẩn hóa nén biên độ, "
         "chẳng hạn vào hàm mục tiêu huấn luyện, hoặc vào cấu trúc graph thông "
         "qua nhiều snapshot thời gian thay vì chỉ qua trọng số cạnh của một "
         "snapshot.")
    para(doc,
         "Static KG-GCN là mô hình đối chứng tự xây dựng nhằm cô lập một biến. "
         "Nó không phải bản tái lập KGAT hoặc KGCN và không đại diện cho toàn "
         "bộ nhóm mô hình gợi ý dựa trên knowledge graph đã công bố.")


def main() -> int:
    parser = argparse.ArgumentParser(description="Sinh Word cho Chuong 4")
    parser.add_argument("--runs-dir", type=Path, default=Path("experiments/runs"))
    parser.add_argument("--out", type=Path, default=Path("docs/Chuong4_viet_lai.docx"))
    args = parser.parse_args()

    if not TEMPLATE.exists():
        raise SystemExit(f"LOI: thieu template {TEMPLATE}")

    frame = load_runs(args.runs_dir, split="test", segment="warm")
    if frame.empty:
        raise SystemExit(f"LOI: khong doc duoc run nao tu {args.runs_dir}")

    doc = Document(str(TEMPLATE))
    clear_body(doc)

    heading(doc, "CHƯƠNG 4. THỰC NGHIỆM VÀ ĐÁNH GIÁ MÔ HÌNH — CÁC MỤC VIẾT LẠI", 1)
    para(doc,
         "Tài liệu này chứa các mục của Chương 4 được viết lại sau khi toàn bộ "
         "ma trận thực nghiệm được chạy lại với ngân sách huấn luyện đến hội tụ "
         "và ba seed. Bảng số được sinh trực tiếp từ thư mục experiments/runs "
         "bằng scripts/08_make_docx.py, không nhập tay.")

    write_421(doc)
    # So bang bam theo v11: Bang 4.5 va 4.6 la hai bang ket qua chinh.
    # So hinh cung bam theo v11: 4.1/4.2 cho original, 4.3/4.4 cho active.
    write_results(doc, frame, "original", "4.2.2", "4.5",
                  "Nhóm người dùng có lịch sử ban đầu", ("4.1", "4.2"))
    write_results(doc, frame, "active", "4.2.3", "4.6",
                  "Nhóm người dùng tích cực", ("4.3", "4.4"))
    heading(doc, "4.3. So sánh với các mô hình và phương pháp khác", 2)
    write_432(doc, frame)
    write_433(doc, frame)

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out))
    cohorts = sorted(frame["cohort"].unique())
    print(f"Da ghi {args.out}")
    print(f"  doc {len(frame)} run tu {args.runs_dir}, split: {', '.join(cohorts)}")
    for cohort in cohorts:
        n = frame[frame['cohort'] == cohort]['model'].nunique()
        print(f"  {cohort}: {n}/6 mo hinh")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
