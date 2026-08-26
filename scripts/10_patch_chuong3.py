#!/usr/bin/env python
"""Va Chuong 3 cua v11 cho khop voi code da dung lai, sinh ra v12.

Vi sao la va tai cho chu khong sinh lai
---------------------------------------
Chuong 3 chua 33 doi tuong cong thuc OMML (Word Equation). python-docx khong doc
va khong ghi duoc chung: khi trich ra chung hien thanh o rong. Sinh lai chuong
nay dong nghia xoa sach cong thuc (3.1)-(3.31). Vi vay script chi thay doan van
quanh cong thuc va KHONG BAO GIO cham vao doan chua OMML.

Nguyen tac an toan
------------------
* Moi ban va gan vao mot doan neo (anchor). Neu neo khong khop DUNG MOT LAN,
  script dung ngay. Va nham vi tri con te hon khong va: no am tham.
* Sau khi va, so cong thuc OMML, so bang va so hinh cua v12 phai bang v11 cong
  dung phan co y them. Kiem o cuoi, khong tin vao doc ky.
* Ban goc v11 khong bao gio bi ghi de. Dau ra la mot file rieng.

Van phong bam theo `docs/VAN_PHONG_DE_AN.md`: giu thuat ngu tieng Anh, dau thap
phan la dau CHAM, khong in dam giua doan, khong dau gach ngang tu tu, cau dai
noi bang dau cham phay, trinh bay chu khong thuyet phuc.

Chay:
    python scripts/10_patch_chuong3.py
    python scripts/10_patch_chuong3.py --out docs/De_an_thac_si_v12.docx
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.utils.logging import get_logger  # noqa: E402

log = get_logger(__name__)

SOURCE = Path("docs/De_an_thac_si_v11.docx")
OMML = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


# ── Tim va sua doan ─────────────────────────────────────────────────────


def find(doc: Document, needle: str):
    """Doan DUY NHAT chua `needle`. Dung neu khong co hoac co nhieu hon mot."""
    hits = [p for p in doc.paragraphs if needle in p.text]
    if len(hits) != 1:
        raise SystemExit(
            f"LOI: neo khop {len(hits)} lan, can dung 1.\n  neo: {needle[:80]!r}\n"
            "Ban v11 da doi khac so voi luc doc. Doc lai roi sua neo."
        )
    para = hits[0]
    if para._p.findall(f".//{OMML}oMath"):
        raise SystemExit(f"LOI: doan neo co chua cong thuc, khong duoc ghi de.\n  {needle[:80]!r}")
    return para


def find_caption_and_index(doc: Document, needle: str):
    """Phan biet caption trong than bai voi dong trong DANH MUC BANG.

    Ca hai cho deu co cung mot chuoi. Caption trong than bai la doan co mot BANG
    nam ngay sau no; dong trong danh muc thi khong.
    """
    hits = [p for p in doc.paragraphs if needle in p.text]
    if len(hits) != 2:
        raise SystemExit(f"LOI: cho doi 2 lan xuat hien cua {needle[:50]!r}, gap {len(hits)}.")
    captions = [p for p in hits
                if (nxt := p._p.getnext()) is not None and nxt.tag == qn("w:tbl")]
    if len(captions) != 1:
        raise SystemExit(f"LOI: khong xac dinh duoc caption than bai cua {needle[:50]!r}.")
    caption = captions[0]
    index = next(p for p in hits if p is not caption)
    return caption, index


def retext(para, text: str) -> None:
    """Thay noi dung doan, giu dinh dang cua run dau tien."""
    for run in para.runs[1:]:
        run._r.getparent().remove(run._r)
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def insert_after(doc: Document, anchor, texts: list[str]):
    """Chen cac doan moi ngay sau `anchor`, dung thu tu da cho."""
    previous = anchor
    made = []
    for text in texts:
        para = doc.add_paragraph(text, style="Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        previous._p.addnext(para._p)
        previous = para
        made.append(para)
    return made


def insert_table_after(doc: Document, anchor, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, header):
        cell.text = ""
        cell.paragraphs[0].add_run(text).bold = True
    for row in rows:
        for cell, text in zip(table.add_row().cells, row):
            cell.text = text
    anchor._p.addnext(table._tbl)
    return table


def cell_containing(doc: Document, needle: str):
    """O bang DUY NHAT chua `needle`."""
    hits = [c for t in doc.tables for r in t.rows for c in r.cells if needle in c.text]
    if len(hits) != 1:
        raise SystemExit(f"LOI: o bang khop {len(hits)} lan, can dung 1: {needle[:60]!r}")
    return hits[0]


# ── Noi dung cac ban va ─────────────────────────────────────────────────


def patch_36_loss(doc: Document) -> None:
    """§3.6: weighted BPR la ablation, khong phai loss cua cohort Original."""
    before = find(doc, "Cần phân biệt hai cấu hình huấn luyện đã tạo ra hai bảng kết quả")
    retext(before,
        "Ngoài BPR chuẩn, đề tài định nghĩa thêm một biến thể có trọng số, trong đó "
        "trọng số dương của cặp visitor-item được dùng khi lấy trung bình loss trên "
        "mini-batch,")

    after = find(doc, "trong đó q(u,i+) là trọng số dương lấy từ cạnh tương tác")
    retext(after,
        "trong đó q(u,i+) là trọng số dương lấy từ cạnh tương tác. Biến thể này chỉ "
        "được dùng trong thí nghiệm loại bỏ thành phần; toàn bộ ma trận thực nghiệm "
        "gồm sáu mô hình, ba seed và hai cohort đều được huấn luyện bằng BPR chuẩn "
        "theo công thức (3.30). Ràng buộc này được cài thành assertion trong lớp cấu "
        "hình, nên một cấu hình đặt loss là weighted_bpr trong ma trận chính sẽ dừng "
        "ngay khi nạp cấu hình thay vì chạy đến cuối rồi cho ra một bảng không so "
        "sánh được. Lý do là trọng số hành vi-thời gian đã tham gia ở bước xây dựng "
        "adjacency và ở phép lan truyền; đưa trọng số vào loss thêm một lần nữa sẽ "
        "làm hai cấu hình khác nhau ở hai biến thay vì một, và mọi chênh lệch đo "
        "được sẽ không quy về được cho riêng behavior-time weighting.")


def patch_34_lambda(doc: Document) -> None:
    """§3.4: lambda=0.01 khong duoc do; bo sung ket qua quet that va mo hinh thu sau."""
    intro = find(doc, "Thiết lập được chọn để báo cáo dùng λ=0.01")
    retext(intro,
        "Đề tài báo cáo hai thiết lập của hệ số suy giảm. Thiết lập thứ nhất dùng "
        "λ=0.01, là giá trị kế thừa từ các công trình multi-behavior trước đó và "
        "không qua bước dò trên dữ liệu của đề tài. Thiết lập thứ hai dùng λ=0.05, "
        "là giá trị được chọn sau khi quét λ trên tập xác thực. Các trọng số hành vi "
        "giữ nguyên ở cả hai thiết lập:")

    caption_35, muc_luc_35 = find_caption_and_index(
        doc, "Bảng 3.5. Trọng số hành vi trong BT-DKGRec-GCN")

    # Cum lambda dat o CUOI §3.4, sau doan chot ve gia dinh cua trong so hanh vi,
    # de khong chen bang moi vao giua Bang 3.5 va doan binh luan cua chinh no.
    closing = find(doc, "Công thức trên thể hiện giả định thực nghiệm")

    method, = insert_after(doc, closing, [
        "Giá trị λ được dò bằng cách quét trên tập xác thực, đo bằng AUC của trọng "
        "số cạnh khi dự báo tương tác mục tiêu; phép đo này đánh giá trực tiếp chất "
        "lượng của trọng số trước khi đưa vào huấn luyện, do đó không tiêu tốn ngân "
        "sách GPU cho từng giá trị λ và không đọc tập test."])

    caption_36 = doc.add_paragraph(style="Normal")
    caption_36.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_36.add_run("Bảng 3.6. Kết quả quét hệ số suy giảm λ trên tập xác thực").bold = True
    method._p.addnext(caption_36._p)

    table_36 = insert_table_after(
        doc, caption_36,
        ["λ", "AUC trên tập xác thực (Original)", "AUC trên tập xác thực (Active)"],
        [["0.000", "0.7862", "0.7637"],
         ["0.010", "0.9045", "0.8566"],
         ["0.020", "0.9180", "0.8825"],
         ["0.050", "0.9185", "0.9012"],
         ["0.100", "0.9069", "0.8963"]])

    tail = doc.add_paragraph(
        "Đỉnh nằm tại λ=0.05 trên cả hai cohort, trong khi λ=0.01 nằm dưới đỉnh ở cả "
        "hai. Trên cửa sổ train dài 97 ngày, λ=0.01 làm trọng số của một tương tác "
        "giảm từ 1.00 xuống 0.38, còn λ=0.05 làm trọng số giảm từ 1.00 xuống 0.008; "
        "thiết lập thứ hai vì vậy nhấn mạnh tính gần đây mạnh hơn nhiều. Đề tài giữ "
        "cả hai thiết lập trong bảng kết quả, ký hiệu là BT-DKGRec-GCN (λ=0.01) và "
        "BT-DKGRec-GCN (λ=0.05), để người đọc thấy được cả cấu hình trước và sau khi "
        "dò tham số; hai mô hình khác nhau đúng một tham số là lambda_decay và dùng "
        "chung toàn bộ phần còn lại của cấu hình.", style="Normal")
    tail.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    table_36._tbl.addnext(tail._p)

    entry = doc.add_paragraph(
        "Bảng 3.6. Kết quả quét hệ số suy giảm λ trên tập xác thực",
        style=muc_luc_35.style)
    muc_luc_35._p.addnext(entry._p)


def patch_352_cancellation(doc: Document) -> None:
    """§3.5.2: he qua cua chuan hoa doi xung voi visitor bac 1."""
    anchor = find(doc, "phép chuẩn hóa giúp node có bậc lớn không lấn át hoàn toàn")
    insert_after(doc, anchor, [
        "Phép chuẩn hóa này có một hệ quả cần được nêu trước khi đọc kết quả thực "
        "nghiệm. Với một visitor chỉ có duy nhất một cạnh trong graph train, bậc có "
        "trọng số của node đó bằng đúng trọng số của cạnh đó, nên thừa số W(u,i) xuất "
        "hiện đồng thời ở tử số và ở mẫu số của phép chuẩn hóa và bị triệt tiêu; "
        "embedding thu được không phụ thuộc vào giá trị của W(u,i), và thứ hạng Top-K "
        "sinh ra cho visitor đó trùng với thứ hạng mà một mô hình không dùng "
        "behavior-time weighting sinh ra. Trong tập huấn luyện Original, 79.6% visitor "
        "có đúng một cạnh, do đó behavior-time weighting chỉ có thể tác động lên phần "
        "visitor còn lại. Đây là lý do đề tài báo cáo nhóm warm tách thêm theo bậc của "
        "visitor trong train, và là ràng buộc cần tính đến khi diễn giải mức chênh "
        "lệch giữa BT-DKGRec-GCN và Static KG-GCN ở Chương 4.",
    ])


def patch_38_training(doc: Document) -> None:
    """§3.8: quy trinh thieu early stopping, nhieu seed va curves.csv."""
    retext(find(doc, "Huấn luyện mô hình đối chứng và BT-DKGRec-GCN."),
           "Huấn luyện mô hình đối chứng và BT-DKGRec-GCN trên cùng một ngân sách, "
           "với ba seed 2020, 2021 và 2022.")
    retext(find(doc, "Báo cáo test warm như kết quả cuối."),
           "Báo cáo test warm như kết quả cuối, kèm phân tách theo bậc của visitor "
           "trong tập huấn luyện.")

    anchor = find(doc, "Quy trình huấn luyện và đánh giá được trình bày tại Hình 3.3")
    para = doc.add_paragraph(
        "Ngân sách huấn luyện được xác định bằng early stopping trên tập xác thực "
        "thay vì bằng một số epoch cố định. Mỗi mô hình được phép chạy tối đa 1000 "
        "epoch, được đánh giá trên tập xác thực sau mỗi 5 epoch, và dừng khi chỉ số "
        "giám sát NDCG@20 không cải thiện sau 20 lần đánh giá liên tiếp; riêng cohort "
        "Active dùng ngưỡng 50 lần vì đường xác thực của cohort này không đơn điệu ở "
        "mọi mô hình, nên một ngưỡng nhỏ hơn sẽ dừng một số mô hình sớm hơn các mô "
        "hình khác và làm mất tính công bằng của so sánh. Ngưỡng được đặt ở cấp cohort "
        "chứ không ở cấp mô hình, nên trong cùng một cohort mọi mô hình nhận cùng một "
        "ngân sách. Toàn bộ quá trình huấn luyện được ghi lại theo từng epoch trong "
        "tệp curves.csv của mỗi lần chạy, gồm loss ở mọi epoch và toàn bộ khối chỉ số "
        "xác thực ở các epoch được đánh giá; nhờ vậy có thể kiểm tra lại rằng mô hình "
        "đối chứng đã hội tụ chứ không bị dừng sớm. Kết quả cuối được báo cáo dưới "
        "dạng trung bình và độ lệch chuẩn trên ba seed.", style="Normal")
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    anchor._p.addprevious(para._p)

    last = find(doc, "Không trộn người dùng có lịch sử và người dùng chưa có lịch sử")
    new = doc.add_paragraph(
        "Không chọn seed và không loại bỏ lần chạy: script sinh bảng đọc toàn bộ "
        "thư mục kết quả.", style=last.style)
    new.paragraph_format.left_indent = last.paragraph_format.left_indent
    last._p.addnext(new._p)


def patch_33_guards(doc: Document) -> None:
    """§3.3: cac dieu kien hop le duoc cai thanh assertion; loc property cu the."""
    anchor = find(doc, "vừa bảo đảm graph có cấu trúc nhất quán, vừa ngăn rò rỉ")
    insert_after(doc, anchor, [
        "Các điều kiện trên được cài đặt thành assertion chạy trong pipeline chứ "
        "không chỉ được phát biểu trong tài liệu. Các assertion kiểm tra rằng split "
        "theo thời gian không chồng lấn; ánh xạ định danh chỉ sinh từ train; side "
        "information chỉ nhận bản ghi có timestamp không vượt quá mốc kết thúc train, "
        "và với item có nhiều bản ghi thì bản ghi mới nhất còn hợp lệ được chọn; mỗi "
        "item chỉ nhận một category; chỉ số node nằm trong phạm vi ánh xạ; tập item "
        "ứng viên là tập con của tập item train; mẫu âm chỉ được lấy từ tập item "
        "train; và bước chọn mô hình chỉ đọc chỉ số của tập xác thực. Khi một "
        "assertion không thỏa, pipeline dừng ngay tại bước đó; các assertion không "
        "được tắt để pipeline chạy qua, vì một lần tắt sẽ làm mọi kết quả sinh ra sau "
        "đó không còn kiểm chứng được.",
    ])

    cell = cell_containing(doc, "Property/value quá thưa hoặc không thuộc phạm vi cấu hình")
    cell.text = ("Property/value xuất hiện dưới 5 lần bị loại; số node property=value "
                 "giữ lại tối đa 50.000 node phổ biến nhất; hai trường categoryid và "
                 "available không tạo node property vì đã được biểu diễn bằng quan hệ "
                 "riêng")


def patch_372_app(doc: Document) -> None:
    """§3.7.2: ung dung doc bo tep xuat, khong truy van Neo4j truc tiep."""
    retext(find(doc, "Ứng dụng cho phép chọn người dùng, xem lịch sử tương tác"),
        "Ứng dụng cho phép chọn người dùng, xem lịch sử tương tác, kiểm tra trọng số "
        "cạnh hành vi-thời gian, quan sát đồ thị con Visitor-Item-Category/Property và "
        "đối chiếu các chỉ số thực nghiệm; ứng dụng đọc trực tiếp bộ tệp xuất của đồ "
        "thị cùng tệp kết quả Top-K của lần chạy, và được tổ chức thành ba phần tương "
        "ứng với lớp lưu vết, lớp chiếu và danh sách Top-K.")

    retext(find(doc, "Kho dữ liệu Neo4j có hai chế độ"),
        "Bước xuất dữ liệu sinh ra một bộ tệp CSV theo lược đồ của đồ thị, gồm node "
        "visitor, item, category, property, node sự kiện của lớp lưu vết và cạnh "
        "visitor-item đã tổng hợp của lớp chiếu; bộ tệp có hai phạm vi, một phạm vi "
        "giới hạn ở các visitor được đánh giá dùng cho trình diễn và một phạm vi đầy "
        "đủ dùng cho kiểm tra lược đồ. Bộ tệp này vừa là dữ liệu nạp vào Neo4j vừa là "
        "nguồn đọc trực tiếp của ứng dụng nguyên mẫu, nên phần trình diễn không phụ "
        "thuộc vào một dịch vụ cơ sở dữ liệu đang chạy. Việc chỉ sử dụng phạm vi huấn "
        "luyện giúp phần trình diễn không đưa hành vi thuộc tập xác thực hoặc tập kiểm "
        "thử vào đồ thị dùng để minh họa suy luận. Tính nhất quán giữa hai lớp được "
        "kiểm bằng một assertion so tổng trọng số sự kiện của lớp lưu vết với trọng số "
        "cạnh của lớp chiếu trên từng cặp visitor-item; trên bản xuất của cohort "
        "Original, phép kiểm đối chiếu 1.570.409 cặp sinh từ 2.024.042 sự kiện và sai "
        "khác tương đối lớn nhất bằng 0.")


PATCHES = (
    ("§3.6  weighted BPR chi la ablation", patch_36_loss),
    ("§3.4  lambda: bo sung ket qua quet that + mo hinh thu sau", patch_34_lambda),
    ("§3.5.2 he qua triet tieu trong so o visitor bac 1", patch_352_cancellation),
    ("§3.8  early stopping, ba seed, curves.csv", patch_38_training),
    ("§3.3  assertion chong ro ri + loc property cu the", patch_33_guards),
    ("§3.7.2 ung dung doc bo tep xuat", patch_372_app),
)


def census(doc: Document) -> dict[str, int]:
    body = doc.element.body
    return {
        "cong thuc": len(body.findall(f".//{OMML}oMath")),
        "bang": len(doc.tables),
        "hinh": len(doc.inline_shapes),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--out", type=Path, default=Path("docs/De_an_thac_si_v12.docx"))
    args = parser.parse_args()

    if args.out.resolve() == args.source.resolve():
        raise SystemExit("LOI: khong ghi de ban goc v11.")

    doc = Document(str(args.source))
    before = census(doc)

    for label, patch in PATCHES:
        patch(doc)
        print(f"  da va  {label}")

    after = census(doc)
    expected = {"cong thuc": before["cong thuc"], "bang": before["bang"] + 1,
                "hinh": before["hinh"]}
    if after != expected:
        raise SystemExit(f"LOI: kiem dem khong khop.\n  truoc: {before}\n  "
                         f"sau:   {after}\n  cho:   {expected}")

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out))
    print(f"\nDa ghi {args.out}")
    print(f"  cong thuc giu nguyen: {after['cong thuc']}")
    print(f"  bang: {before['bang']} -> {after['bang']} (them Bang 3.6)")
    print(f"  hinh giu nguyen: {after['hinh']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
