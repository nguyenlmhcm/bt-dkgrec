"""Kiem bo va Chuong 1 va Chuong 2 (`scripts/11_patch_chuong12.py`).

Cung tinh than voi `test_patch_chuong3.py`: rui ro lon nhat khong phai va sai
chu ma la va nham CHO, hoac mot cau da sua roi lai quay ve o ban sau. Vi vay
test o day gac ba thu: hai cong thuc (2.1) va (2.2) con nguyen, cac cau sai da
biet khong duoc xuat hien lai, va ban v12 khong bi ghi de.
"""

from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

import pytest

# Chi VPS moi sinh tai lieu; Colab cai requirements-colab.txt va khong co
# python-docx. Bo qua ca module thay vi lam gay `pytest` cua notebook train.
pytest.importorskip("docx", reason="can python-docx — chi chay tren VPS")

from docx import Document  # noqa: E402

ROOT = Path(__file__).resolve().parents[1]
V12 = ROOT / "docs" / "De_an_thac_si_v12.docx"
OMML = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

pytestmark = pytest.mark.skipif(not V12.is_file(), reason="chua co ban v12")


def load():
    spec = importlib.util.spec_from_file_location(
        "patch_chuong12", ROOT / "scripts" / "11_patch_chuong12.py")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


@pytest.fixture(scope="module")
def patched(tmp_path_factory) -> Path:
    out = tmp_path_factory.mktemp("docx") / "v13.docx"
    module = load()
    argv = sys.argv
    sys.argv = ["11_patch_chuong12.py", "--out", str(out), "--no-highlight"]
    try:
        module.main()
    finally:
        sys.argv = argv
    return out


def all_text(path: Path) -> str:
    doc = Document(str(path))
    body = "\n".join(p.text for p in doc.paragraphs)
    cells = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    return body + "\n" + cells


def than_bai(path: Path) -> str:
    """Chi cac doan van, khong lay o bang.

    Bang danh muc sua doi TRICH LAI nguyen van cac cau da bo — do la viec cua
    no. Neu quet ca bang thi moi kiem tra hoi quy deu bao dong gia.
    """
    return "\n".join(p.text for p in Document(str(path)).paragraphs)


def doan_chuong_1_2(path: Path) -> list[str]:
    """Cac doan tu 'CHUONG 1' den truoc 'CHUONG 3' — dung pham vi ban va nay."""
    out: list[str] = []
    trong = False
    for para in Document(str(path)).paragraphs:
        text = para.text.strip()
        if text.startswith("CHƯƠNG 1"):
            trong = True
        elif text.startswith("CHƯƠNG 3"):
            break
        if trong and text:
            out.append(text)
    return out


def census(path: Path) -> dict:
    doc = Document(str(path))
    return {
        "cong thuc": len(doc.element.body.findall(f".//{OMML}oMath")),
        "bang": len(doc.tables),
        "hinh": len(doc.inline_shapes),
    }


# ── Khong duoc lam hong tai lieu ────────────────────────────────────────


def test_giu_nguyen_cong_thuc_bang_hinh(patched):
    """Ban va chi thay doan van; khong duoc them bot doi tuong nao."""
    assert census(patched) == census(V12)


def test_khong_ghi_de_ban_v12(patched):
    assert V12.is_file()
    assert patched.resolve() != V12.resolve()


def test_hai_cong_thuc_chuong_2_con_nguyen(patched):
    """(2.1) va (2.2) la ly do phai va thay vi sinh lai ca chuong."""
    text = all_text(patched)
    assert "(2.1)" in text
    assert "(2.2)" in text


# ── Cac cau sai da biet khong duoc quay lai ─────────────────────────────


@pytest.mark.parametrize("cau_cu", [
    # §2.6 bao cao ket qua cua v11, da bi ban dung lai bac bo
    "vẫn thấp hơn nhẹ",
    "gần LightGCN nhưng",
    # §2.7 cau cut, khong co menh de chinh
    "Thay vì trình bày lại các công thức chuẩn",
    # loi nhan doi tu mot lan tim-thay the thuat ngu
    "(người dùng có lịch sử)",
    # §2.6 cum vo nghia
    "cửa sổ tài liệu chính",
    # §2.4 cau tu khuyen nghi
    "nên được đặt ở phạm vi hẹp",
])
def test_cau_sai_khong_con(patched, cau_cu):
    assert cau_cu not in than_bai(patched)


# ── Noi dung moi phai co mat ────────────────────────────────────────────


@pytest.mark.parametrize("moc", [
    "1.4.1. Đối tượng nghiên cứu",
    "1.4.2. Phạm vi nghiên cứu",
    "1.5. Nội dung và phương pháp nghiên cứu",
    "gồm sáu bậc",
    "thoả đồng thời hai điều kiện",
    "MBGCN mô hình hóa nhiều loại hành vi",
    "[19] B. Jin",
    "[20] L. Xia",
])
def test_noi_dung_moi_co_mat(patched, moc):
    assert moc in all_text(patched)


def test_thu_muc_tham_khao_du_hai_muc_moi(patched):
    """Khong duoc trich [19]/[20] trong than bai ma quen them vao thu muc."""
    text = all_text(patched)
    for so in ("[19]", "[20]"):
        assert text.count(so) >= 2, f"{so} phai xuat hien ca o than bai lan thu muc"


def test_danh_muc_sua_doi_duoc_noi_vao_bang_cu(patched):
    """Mot bang danh muc sua doi duy nhat, khong sinh them bang thu hai."""
    doc = Document(str(patched))
    bang = [t for t in doc.tables
            if t.rows and "Mục" in t.rows[0].cells[0].text and len(t.columns) == 3]
    assert len(bang) == 1
    dong_v13 = [r for r in bang[0].rows if r.cells[0].text.startswith("[v13]")]
    assert len(dong_v13) >= 10


# ── Van phong ───────────────────────────────────────────────────────────


def test_khong_xung_ho_truc_tiep(patched):
    """Register eng-impersonal: tai lieu mo ta he thong, khong goi nguoi doc.

    Chi quet Chuong 1 va Chuong 2. Loi cam on co "gia dinh, ban be", trong do
    "ban" la danh tu chu khong phai dai tu xung ho.
    """
    for text in doan_chuong_1_2(patched):
        assert " bạn " not in f" {text} "
        assert not text.startswith("Bạn ")
