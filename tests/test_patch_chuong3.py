"""Kiem bo va Chuong 3 (`scripts/10_patch_chuong3.py`).

Rui ro lon nhat cua viec va file .docx khong phai la va sai chu, ma la va nham
CHO: Word van mo binh thuong, cong thuc van hien, va khong ai biet mot doan da
bi ghi de nham. Vi vay test o day tap trung vao ba dieu: neo phai duy nhat, so
cong thuc phai giu nguyen, va ban goc v11 khong duoc dong toi.
"""

from __future__ import annotations

import importlib.util
import shutil
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
V11 = ROOT / "docs" / "De_an_thac_si_v11.docx"
OMML = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

pytestmark = pytest.mark.skipif(not V11.is_file(), reason="chua co ban v11")


def load():
    spec = importlib.util.spec_from_file_location(
        "patch_chuong3", ROOT / "scripts" / "10_patch_chuong3.py")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def run_patch(out: Path, *extra: str) -> Path:
    import sys
    module = load()
    argv = sys.argv
    sys.argv = ["10_patch_chuong3.py", "--source", str(V11), "--out", str(out), *extra]
    try:
        assert module.main() == 0
    finally:
        sys.argv = argv
    return out


@pytest.fixture(scope="module")
def patched(tmp_path_factory) -> Path:
    """Ban ra soat: co to nen va co danh muc sua doi."""
    return run_patch(tmp_path_factory.mktemp("v12") / "v12.docx")


@pytest.fixture(scope="module")
def clean(tmp_path_factory) -> Path:
    """Ban nop: khong to nen, khong danh muc."""
    return run_patch(tmp_path_factory.mktemp("v12c") / "v12_clean.docx", "--no-highlight")


def count_omath(path: Path) -> int:
    return len(Document(str(path)).element.body.findall(f".//{OMML}oMath"))


def test_khong_mat_cong_thuc(patched):
    """Diem chet cua ca cach lam: python-docx khong tai tao duoc OMML."""
    assert count_omath(patched) == count_omath(V11)


def test_ban_nop_them_dung_mot_bang(clean):
    """Ban sach chi them Bang 3.6."""
    assert len(Document(str(clean)).tables) == len(Document(str(V11)).tables) + 1


def test_ban_ra_soat_them_hai_bang(patched):
    """Ban ra soat them ca bang danh muc sua doi."""
    assert len(Document(str(patched)).tables) == len(Document(str(V11)).tables) + 2


def test_giu_nguyen_so_hinh(patched):
    assert len(Document(str(patched)).inline_shapes) == len(Document(str(V11)).inline_shapes)


def test_ban_goc_khong_bi_sua(patched):
    """v11 la dac ta goc; mot lan ghi de la mat vinh vien."""
    doc = Document(str(V11))
    assert not [p for p in doc.paragraphs if "Biến thể này chỉ được dùng" in p.text]
    assert [p for p in doc.paragraphs if "Thiết lập được chọn để báo cáo dùng λ=0.01" in p.text]


def test_tu_choi_ghi_de_ban_goc(tmp_path):
    module = load()
    import sys
    argv = sys.argv
    sys.argv = ["10_patch_chuong3.py", "--source", str(V11), "--out", str(V11)]
    try:
        with pytest.raises(SystemExit):
            module.main()
    finally:
        sys.argv = argv


def test_neo_khong_duy_nhat_thi_dung(tmp_path):
    """`find()` phai nem khi mot chuoi xuat hien hai lan, khong duoc lay cai dau."""
    module = load()
    doc = Document(str(V11))
    with pytest.raises(SystemExit):
        module.find(doc, "Bảng 3.5. Trọng số hành vi trong BT-DKGRec-GCN")


def test_khong_ghi_de_doan_chua_cong_thuc(tmp_path):
    module = load()
    doc = Document(str(V11))
    omath_paras = [p for p in doc.paragraphs if p._p.findall(f".//{OMML}oMath")]
    if not omath_paras:
        pytest.skip("v11 khong co cong thuc trong doan van")
    needle = omath_paras[0].text[:20]
    if not needle.strip():
        pytest.skip("doan cong thuc khong co van ban de neo")
    with pytest.raises(SystemExit):
        module.find(doc, needle)


@pytest.mark.parametrize("cau", [
    "Biến thể này chỉ được dùng trong thí nghiệm loại bỏ thành phần",
    "Thiết lập thứ hai dùng λ=0.05",
    "Phép chuẩn hóa này có một hệ quả",
    "Ngân sách huấn luyện được xác định bằng early stopping",
    "Các điều kiện trên được cài đặt thành assertion",
    "Bước xuất dữ liệu sinh ra một bộ tệp CSV",
])
def test_noi_dung_moi_co_mat(patched, cau):
    assert [p for p in Document(str(patched)).paragraphs if cau in p.text]


@pytest.mark.parametrize("cau_cu", [
    "Các giá trị này là cấu hình thực nghiệm được lựa chọn trên tập xác thực",
    "Original người dùng có lịch sử đánh giá dùng weighted BPR",
    "Kho dữ liệu Neo4j có hai chế độ",
])
def test_cau_sai_da_bi_go(patched, cau_cu):
    assert not [p for p in Document(str(patched)).paragraphs if cau_cu in p.text]


def test_dau_thap_phan_la_dau_cham(patched):
    """Quy tac 2 cua docs/VAN_PHONG_DE_AN.md."""
    doc = Document(str(patched))
    moi = [p.text for p in doc.paragraphs if "λ=0.05" in p.text or "79.6%" in p.text]
    assert moi, "khong tim thay doan moi de kiem"
    for text in moi:
        assert "0,05" not in text and "79,6" not in text


# ── Danh dau cho sua ────────────────────────────────────────────────────


def test_danh_muc_sua_doi_nam_o_dau_tai_lieu(patched):
    """Trang bia v11 la mot BANG, nen de sai cho la danh muc bi day xuong duoi."""
    doc = Document(str(patched))
    header = [c.text.strip() for c in doc.tables[0].rows[0].cells]
    assert header == ["Mục", "Nội dung sửa", "Vì sao"]
    assert doc.paragraphs[0].text.strip() == "DANH MỤC SỬA ĐỔI SO VỚI BẢN v11"


def test_danh_muc_liet_ke_du_sau_muc_da_va(patched):
    muc = {r.cells[0].text.strip() for r in Document(str(patched)).tables[0].rows[1:]}
    assert muc == {"§3.6", "§3.4", "§3.5.2", "§3.8", "§3.3.3", "Bảng 3.4", "§3.7.2"}


def test_cho_sua_duoc_to_nen(patched):
    """Khong to nen thi khong ai tim ra cho sua giua 16.606 tu."""
    from docx.enum.text import WD_COLOR_INDEX
    doc = Document(str(patched))
    vang = [p.text for p in doc.paragraphs
            for r in p.runs if r.font.highlight_color == WD_COLOR_INDEX.YELLOW]
    assert any("Biến thể này chỉ được dùng" in t for t in vang)
    assert any("Phép chuẩn hóa này có một hệ quả" in t for t in vang)
    assert any("Ngân sách huấn luyện được xác định" in t for t in vang)


def test_ban_nop_khong_con_dau_vet_ra_soat(clean):
    from docx.enum.text import WD_COLOR_INDEX
    doc = Document(str(clean))
    assert not [r for p in doc.paragraphs for r in p.runs
                if r.font.highlight_color == WD_COLOR_INDEX.YELLOW]
    assert "DANH MỤC SỬA ĐỔI" not in doc.paragraphs[0].text
    assert [p for p in doc.paragraphs if "Biến thể này chỉ được dùng" in p.text]
